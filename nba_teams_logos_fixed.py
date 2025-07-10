import os
import json
import sys
import requests
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import traceback
import urllib.parse
import io
import logging  # Add logging import

# Enable urllib3 debugging
import urllib3
urllib3.add_stderr_logger(level=logging.DEBUG)

# A régi nba_logos_to_word.py tartalma javítva

# Próbáljuk importálni a CairoSVG-t az SVG konvertáláshoz
try:
    import cairosvg
    CAIROSVG_AVAILABLE = True
    print("CairoSVG sikeresen betöltve. Az SVG-k konvertálása lehetséges.")
except ImportError:
    CAIROSVG_AVAILABLE = False
    print("CairoSVG nem elérhető. Az SVG képek nem lesznek konvertálva PNG-re.")

# Az NBA csapatok definíciója a csapatkódokkal
NBA_TEAMS = [
    {"id": 1610612738, "abbreviation": "BOS", "name": "Celtics", "city": "Boston"},
    {"id": 1610612751, "abbreviation": "BKN", "name": "Nets", "city": "Brooklyn"},
    {"id": 1610612752, "abbreviation": "NYK", "name": "Knicks", "city": "New York"},
    {"id": 1610612755, "abbreviation": "PHI", "name": "76ers", "city": "Philadelphia"},
    {"id": 1610612761, "abbreviation": "TOR", "name": "Raptors", "city": "Toronto"},
    {"id": 1610612741, "abbreviation": "CHI", "name": "Bulls", "city": "Chicago"},
    {"id": 1610612739, "abbreviation": "CLE", "name": "Cavaliers", "city": "Cleveland"},
    {"id": 1610612765, "abbreviation": "DET", "name": "Pistons", "city": "Detroit"},
    {"id": 1610612754, "abbreviation": "IND", "name": "Pacers", "city": "Indiana"},
    {"id": 1610612749, "abbreviation": "MIL", "name": "Bucks", "city": "Milwaukee"},
    {"id": 1610612737, "abbreviation": "ATL", "name": "Hawks", "city": "Atlanta"},
    {"id": 1610612766, "abbreviation": "CHA", "name": "Hornets", "city": "Charlotte"},
    {"id": 1610612748, "abbreviation": "MIA", "name": "Heat", "city": "Miami"},
    {"id": 1610612753, "abbreviation": "ORL", "name": "Magic", "city": "Orlando"},
    {"id": 1610612764, "abbreviation": "WAS", "name": "Wizards", "city": "Washington"},
    {"id": 1610612743, "abbreviation": "DEN", "name": "Nuggets", "city": "Denver"},
    {"id": 1610612750, "abbreviation": "MIN", "name": "Timberwolves", "city": "Minnesota"},
    {"id": 1610612760, "abbreviation": "OKC", "name": "Thunder", "city": "Oklahoma City"},
    {"id": 1610612757, "abbreviation": "POR", "name": "Trail Blazers", "city": "Portland"},
    {"id": 1610612762, "abbreviation": "UTA", "name": "Jazz", "city": "Utah"},
    {"id": 1610612744, "abbreviation": "GSW", "name": "Warriors", "city": "Golden State"},
    {"id": 1610612746, "abbreviation": "LAC", "name": "Clippers", "city": "Los Angeles"},
    {"id": 1610612747, "abbreviation": "LAL", "name": "Lakers", "city": "Los Angeles"},
    {"id": 1610612756, "abbreviation": "PHX", "name": "Suns", "city": "Phoenix"},
    {"id": 1610612758, "abbreviation": "SAC", "name": "Kings", "city": "Sacramento"},
    {"id": 1610612742, "abbreviation": "DAL", "name": "Mavericks", "city": "Dallas"},
    {"id": 1610612745, "abbreviation": "HOU", "name": "Rockets", "city": "Houston"},
    {"id": 1610612763, "abbreviation": "MEM", "name": "Grizzlies", "city": "Memphis"},
    {"id": 1610612740, "abbreviation": "NOP", "name": "Pelicans", "city": "New Orleans"},
    {"id": 1610612759, "abbreviation": "SAS", "name": "Spurs", "city": "San Antonio"}
]

def convert_svg_to_png(svg_data, output_path):
    """
    SVG adat konvertálása PNG-re és mentése
    """
    try:
        if CAIROSVG_AVAILABLE:
            # CairoSVG használata
            if isinstance(svg_data, str):  # Ha fájl útvonal
                cairosvg.svg2png(url=svg_data, write_to=output_path)
            else:  # Ha bytes
                cairosvg.svg2png(bytestring=svg_data, write_to=output_path)
            return True
        else:
            # Ha nincs CairoSVG, akkor mentjük előbb az SVG-t, majd az inkscape-et használjuk
            if isinstance(svg_data, bytes):
                svg_path = output_path.replace('.png', '.svg')
                with open(svg_path, 'wb') as f:
                    f.write(svg_data)
                
                # Inkscape parancssoros hívása
                inkscape_cmd = f"inkscape --export-filename={output_path} {svg_path}"
                if os.system(inkscape_cmd) == 0 and os.path.exists(output_path):
                    return True
    except Exception as e:
        print(f"Hiba az SVG->PNG konvertálás során: {str(e)}")
    
    return False

def create_text_image(text, output_path, width=200, height=200):
    """
    Egyszerű szöveges kép létrehozása
    """
    try:
        # Létrehozunk egy üres képet
        img = PILImage.new('RGB', (width, height), color=(240, 240, 240))
        
        # Szöveg kirajzolásához betöltjük az ImageDraw modult
        from PIL import ImageDraw, ImageFont
        draw = ImageDraw.Draw(img)
        
        # Alapértelmezett betűtípus
        font = ImageFont.load_default()
        
        # Szöveg elhelyezése középre
        lines = text.split("\n")
        y_pos = (height - len(lines) * 20) // 2
        
        for line in lines:
            # Szöveghossz becslése (pontos mérés nem mindig elérhető)
            textwidth = font.getlength(line) if hasattr(font, 'getlength') else len(line) * 6
            x_pos = (width - textwidth) // 2
            
            # Szöveg kirajzolása
            draw.text((x_pos, y_pos), line, fill=(0, 0, 0))
            y_pos += 20
        
        # Kép mentése
        img.save(output_path)
        return True
    except Exception as e:
        print(f"Hiba a szöveges kép létrehozása közben: {str(e)}")
        return False

def get_logo_for_team(team_id, team_abbreviation, team_name, team_city, output_dir):
    """
    Adott csapathoz logó letöltése és visszaadása
    """
    # Meghatározzuk a kimeneti fájl nevét
    png_filename = os.path.join(output_dir, f"{team_id}_{team_abbreviation}_{team_name}.png")
    
    # Lehetséges források a megfelelő sorrendben
    logo_sources = [
        # Wikimedia sources - SVG
        {
            "url": f"https://upload.wikimedia.org/wikipedia/en/logo/0/0d/National_Basketball_Association_{team_abbreviation}.svg",
            "type": "svg"
        },
        # Wikimedia alternative source
        {
            "url": f"https://upload.wikimedia.org/wikipedia/en/0/0d/{team_city.replace(' ', '_')}_{team_name.replace(' ', '_')}_logo.svg",
            "type": "svg"
        },
        # NBA API sources
        {
            "url": f"https://cdn.nba.com/logos/nba/{team_id}/global/L/logo.svg",
            "type": "svg"
        },
        {
            "url": f"https://cdn.nba.com/logos/nba/{team_id}/primary/L/logo.svg",
            "type": "svg"
        },
        {
            "url": f"https://cdn.nba.com/logos/nba/{team_id}/global/D/logo.svg",
            "type": "svg"
        },
        # Direct specific mappings for problematic teams
        # Warriors
        {
            "url": "https://upload.wikimedia.org/wikipedia/en/0/01/Golden_State_Warriors_logo.svg",
            "type": "svg",
            "for_team": "GSW"
        },
        # Lakers
        {
            "url": "https://upload.wikimedia.org/wikipedia/commons/3/3c/Los_Angeles_Lakers_logo.svg",
            "type": "svg",
            "for_team": "LAL"
        },
        # Clippers
        {
            "url": "https://upload.wikimedia.org/wikipedia/en/b/bb/Los_Angeles_Clippers_%282015%29.svg",
            "type": "svg",
            "for_team": "LAC"
        },
        # Suns
        {
            "url": "https://upload.wikimedia.org/wikipedia/en/d/dc/Phoenix_Suns_logo.svg",
            "type": "svg",
            "for_team": "PHX"
        },
        # Mavericks
        {
            "url": "https://upload.wikimedia.org/wikipedia/en/9/97/Dallas_Mavericks_logo.svg",
            "type": "svg",
            "for_team": "DAL"
        },
        # Rockets
        {
            "url": "https://upload.wikimedia.org/wikipedia/en/2/28/Houston_Rockets.svg",
            "type": "svg",
            "for_team": "HOU"
        },
        # Grizzlies
        {
            "url": "https://upload.wikimedia.org/wikipedia/en/f/f1/Memphis_Grizzlies.svg",
            "type": "svg",
            "for_team": "MEM"
        },
        # Pelicans
        {
            "url": "https://upload.wikimedia.org/wikipedia/en/0/0d/New_Orleans_Pelicans_logo.svg",
            "type": "svg",
            "for_team": "NOP"
        },
        # Spurs
        {
            "url": "https://upload.wikimedia.org/wikipedia/en/a/a2/San_Antonio_Spurs.svg",
            "type": "svg",
            "for_team": "SAS"
        },
        # PNG sources
        {
            "url": f"https://cdn.nba.com/logos/nba/{team_id}/global/L/logo.png",
            "type": "png"
        },
        {
            "url": f"https://cdn.nba.com/logos/nba/{team_id}/primary/L/logo.png",
            "type": "png"
        },
        # Backup option - ESPN source
        {
            "url": f"https://a.espncdn.com/i/teamlogos/nba/500/scoreboard/{team_abbreviation.lower()}.png",
            "type": "png"
        }
    ]
    
    # Próbáljuk meg letölteni a logót a különböző forrásokból
    for source in logo_sources:
        # Ha a forrás egy adott csapathoz van rendelve és nem ez a csapat, lépjünk tovább
        if "for_team" in source and source["for_team"] != team_abbreviation:
            continue
        
        url = source["url"]
        
        try:
            # Debug: print the exact URL we're requesting
            print(f"Attempting to download from URL: {url}")
            
            response = requests.get(url, stream=True)
            if response.status_code == 200:
                # Sikeres letöltés
                content_type = response.headers.get("Content-Type", "")
                
                if "svg" in content_type or source["type"] == "svg":
                    # SVG fájl konvertálása PNG-re
                    svg_data = response.content
                    if convert_svg_to_png(svg_data, png_filename):
                        print(f"Logó letöltve és konvertálva: {team_city} {team_name} ({url})")
                        return png_filename
                else:
                    # Közvetlenül PNG vagy más formátumú kép
                    with open(png_filename, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    
                    # Ellenőrizzük, hogy érvényes kép-e
                    try:
                        with PILImage.open(png_filename) as img:
                            if img.width > 50 and img.height > 50:  # Minimális méret
                                print(f"Logó letöltve: {team_city} {team_name} ({url})")
                                return png_filename
                            else:
                                print(f"Túl kicsi a letöltött logó, próbáljuk a következő forrást...")
                    except Exception as e:
                        print(f"A letöltött fájl nem érvényes kép: {str(e)}")
                        # Töröljük a hibás fájlt
                        if os.path.exists(png_filename):
                            os.remove(png_filename)
        except Exception as e:
            print(f"Hiba a logó letöltése közben ({url}): {str(e)}")
    
    # Ha minden forrás sikertelen volt, készítsünk egy szöveges helyettesítő képet
    print(f"Nem sikerült letölteni a logót: {team_city} {team_name}, helyettesítő kép készítése")
    create_text_image(f"{team_city}\n{team_name}", png_filename)
    
    return png_filename

def download_nba_logos(output_dir="nba_logos_fixed"):
    """
    NBA csapatok logóinak letöltése és helyettesítő képek készítése szükség esetén
    """
    # Létrehozzuk a kimeneti mappát
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    print(f"NBA logók letöltése a következő mappába: {output_dir}")
    
    # Eredmény követése
    download_results = {"successful": [], "failed": []}
    
    # Logók letöltése minden csapathoz
    for team in NBA_TEAMS:
        team_id = team["id"]
        team_abbreviation = team["abbreviation"]
        team_name = team["name"]
        team_city = team["city"]
        
        # Logó letöltése
        logo_path = get_logo_for_team(team_id, team_abbreviation, team_name, team_city, output_dir)
        
        # Ellenőrizzük, hogy jó-e a logó
        success = False
        try:
            with PILImage.open(logo_path) as img:
                if img.width > 50 and img.height > 50:
                    success = True
        except Exception:
            success = False
        
        if success:
            download_results["successful"].append({
                "team_id": team_id,
                "name": f"{team_city} {team_name}",
                "abbreviation": team_abbreviation,
                "logo_path": os.path.basename(logo_path)
            })
        else:
            download_results["failed"].append({
                "team_id": team_id,
                "name": f"{team_city} {team_name}",
                "abbreviation": team_abbreviation
            })
    
    # Mentsük el az eredményeket JSON-be
    results_file = os.path.join(output_dir, "download_results.json")
    with open(results_file, 'w', encoding='utf-8') as f:
        json.dump(download_results, f, indent=2, ensure_ascii=False)
    
    print(f"\nÖsszegzés:")
    print(f"Sikeresen letöltve: {len(download_results['successful'])} logó")
    print(f"Sikertelen: {len(download_results['failed'])} csapat")
    
    if download_results["failed"]:
        print(f"Sikertelen letöltések: {[t['name'] for t in download_results['failed']]}")
    
    return len(download_results["successful"])

def create_nba_teams_word_document(logos_dir="nba_logos_fixed", output_file="nba_teams_logos.docx"):
    """
    Word dokumentum létrehozása NBA csapatok logóival
    """
    print(f"Word dokumentum létrehozása a következő mappából: {logos_dir}")
    
    # Ellenőrizzük, hogy létezik-e a mappa
    if not os.path.exists(logos_dir):
        print(f"A megadott mappa nem létezik: {logos_dir}")
        return False
    
    # Gyűjtsük össze az elérhető PNG/JPG képeket
    logo_files = []
    for root, _, files in os.walk(logos_dir):
        for file in files:
            if file.lower().endswith(('.png', '.jpg', '.jpeg')):
                logo_files.append(os.path.join(root, file))
    
    if not logo_files:
        print(f"Nem találtunk képfájlokat a következő mappában: {logos_dir}")
        return False
    
    print(f"Összesen {len(logo_files)} képet találtunk")
    
    # Hozzuk létre az új Word dokumentumot
    doc = Document()
    
    # Állítsuk be a dokumentum tulajdonságait
    doc.core_properties.title = "NBA Csapatok Logói"
    doc.core_properties.author = "NBA Logo Downloader"
    
    # Címlap
    title = doc.add_heading('NBA Csapatok', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Rövid magyarázat
    p = doc.add_paragraph('Az alábbi dokumentum tartalmazza az NBA csapatok logóit és városaikat.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Gyűjtsük ki a csapatinformációkat a fájlnevekből és az NBA_TEAMS adatokból
    teams_info = {}
    
    # Először az NBA_TEAMS adatokat használjuk alapként
    for team in NBA_TEAMS:
        team_id = team["id"]
        teams_info[team_id] = {
            'name': team["name"],
            'abbreviation': team["abbreviation"],
            'city': team["city"],
            'logo_path': None  # Először nincs logó, majd a fájlnevek alapján beállítjuk
        }
    
    # Rendeljük hozzá a logó fájlokat a megfelelő csapatokhoz
    for logo_path in logo_files:
        file_name = os.path.basename(logo_path)
        # Várható formátum: ID_RÖVIDÍTÉS_NÉV.png vagy ID_RÖVIDÍTÉS_NÉV.svg
        for team_id, team_info in teams_info.items():
            abbr = team_info['abbreviation']
            if f"{team_id}_{abbr}_" in file_name:
                teams_info[team_id]['logo_path'] = logo_path
                break
    
    # Ha még mindig vannak csapatok logó nélkül, készítsünk helyettesítő képeket
    for team_id, team_info in teams_info.items():
        if not team_info['logo_path']:
            team_name = team_info['name']
            team_abbr = team_info['abbreviation']
            team_city = team_info['city']
            
            print(f"Hiányzó logó pótlása: {team_city} {team_name}")
            placeholder_png = os.path.join(logos_dir, f"{team_id}_{team_abbr}_{team_name}.png")
            create_text_image(f"{team_city}\n{team_name}", placeholder_png)
            teams_info[team_id]['logo_path'] = placeholder_png
    
    # Rendezzük a csapatokat konferenciánként (Keleti és Nyugati)
    eastern_teams = []
    western_teams = []
    
    # Keleti konferencia csapatai
    eastern_abbrs = ["BOS", "BKN", "NYK", "PHI", "TOR", "CHI", "CLE", "DET", "IND", "MIL", 
                      "ATL", "CHA", "MIA", "ORL", "WAS"]
    
    # Csapatok rendezése konferenciák szerint
    for team_id, team_info in teams_info.items():
        if team_info['abbreviation'] in eastern_abbrs:
            eastern_teams.append((team_id, team_info))
        else:
            western_teams.append((team_id, team_info))
    
    # Ellenőrizzük, hogy minden csapat szerepel-e
    all_abbrs = set([team_info['abbreviation'] for team_id, team_info in eastern_teams + western_teams])
    expected_abbrs = set([team["abbreviation"] for team in NBA_TEAMS])
    if all_abbrs != expected_abbrs:
        missing = expected_abbrs - all_abbrs
        if missing:
            print(f"FIGYELEM: Hiányzó csapatok: {missing}")
    
    # Dokumentum szélességének beállítása (hogy elférjen 4 oszlop)
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4 szélesség
    section.page_height = Cm(29.7)  # A4 magasság
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    
    # Minden csapat egy oldalon, 4 oszlopos elrendezésben
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    
    # Állítsuk be a táblázat oszlopainak szélességét
    for cell in table.columns[0].cells:
        cell.width = Cm(4.5)
    for cell in table.columns[1].cells:
        cell.width = Cm(4.5)
    for cell in table.columns[2].cells:
        cell.width = Cm(4.5)
    for cell in table.columns[3].cells:
        cell.width = Cm(4.5)
    
    # Minden csapat egy listába
    all_teams = []
    
    # Az Eastern Conference csapatai jönnek előre
    eastern_teams.sort(key=lambda x: x[1]['city'])  # Rendezzük városnév szerint
    for team in eastern_teams:
        all_teams.append(team)
    
    # Utána a Western Conference csapatai
    western_teams.sort(key=lambda x: x[1]['city'])  # Rendezzük városnév szerint
    for team in western_teams:
        all_teams.append(team)
    
    # Feltöltjük a táblázatot minden csapattal
    current_row = None
    
    for i, (team_id, team_info) in enumerate(all_teams):
        if i % 4 == 0:
            current_row = table.add_row()
            current_row.height = Cm(2.5)  # Kisebb sormagasság a 8 sor érdekében
        
        cell = current_row.cells[i % 4]
        team_name = team_info.get('name', '')
        team_city = team_info.get('city', '')
        team_abbr = team_info.get('abbreviation', '')
        logo_path = team_info.get('logo_path')
        
        if logo_path and os.path.exists(logo_path):
            # Hozzáadjuk a képet és a nevet
            p = cell.paragraphs[0]
            run = p.add_run()
            
            try:
                # Beillesztjük a képet, még kisebb méretben a 8 sor érdekében
                run.add_picture(logo_path, width=Cm(2.0))
            except Exception as e:
                print(f"Hiba a kép beillesztésekor ({team_name}): {str(e)}")
                run.add_text(f"[{team_city} {team_name}]")
            
            # Adjuk hozzá a csapat nevét és városát
            name_text = f"{team_city} {team_name}"
            
            # Név hozzáadása a cellához
            name_paragraph = cell.add_paragraph(name_text)
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Beállítjuk a szöveg betűméretét 9-re
            for run in name_paragraph.runs:
                run.font.size = Pt(9)
            
            # Csökkentsük a bekezdések közötti térközt
            name_paragraph.paragraph_format.space_after = Pt(2)
            
        else:
            # Ha nincs kép, csak a nevet és várost adjuk hozzá
            p = cell.paragraphs[0]
            p.text = f"{team_city} {team_name}"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Beállítjuk a szöveg betűméretét 9-re
            for run in p.runs:
                run.font.size = Pt(9)
    
    # Mentés
    try:
        doc.save(output_file)
        print(f"A Word dokumentum sikeresen létrehozva: {output_file}")
        return True
    except Exception as e:
        print(f"Hiba a Word dokumentum mentésekor: {str(e)}")
        traceback.print_exc()
        return False

def main():
    # Beállítjuk a mappát, ahol a logókat tároljuk
    logos_dir = "nba_logos_fixed"
    output_file = "nba_teams_logos.docx"
    
    # Letöltjük a logókat
    print("Logók letöltése...")
    download_nba_logos(logos_dir)
    
    # Word dokumentum létrehozása
    print("\nWord dokumentum készítése...")
    result = create_nba_teams_word_document(logos_dir, output_file)
    
    if result:
        print(f"A dokumentum elkészült: {output_file}")
    else:
        print("A dokumentum létrehozása sikertelen.")

if __name__ == "__main__":
    main() 