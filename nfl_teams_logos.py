import os
import json
import sys
import requests
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import traceback
import urllib.parse
import io

# Próbáljuk importálni a CairoSVG-t az SVG konvertáláshoz
try:
    import cairosvg
    CAIROSVG_AVAILABLE = True
    print("CairoSVG sikeresen betöltve. Az SVG-k konvertálása lehetséges.")
except ImportError:
    CAIROSVG_AVAILABLE = False
    print("CairoSVG nem elérhető. Az SVG képek nem lesznek konvertálva PNG-re.")

# Az NFL csapatok definíciója a csapatkódokkal és divíziókkal
NFL_TEAMS = [
    # AFC East
    {"id": "BUF", "name": "Bills", "city": "Buffalo", "division": "AFC East"},
    {"id": "MIA", "name": "Dolphins", "city": "Miami", "division": "AFC East"},
    {"id": "NE", "name": "Patriots", "city": "New England", "division": "AFC East"},
    {"id": "NYJ", "name": "Jets", "city": "New York", "division": "AFC East"},
    
    # AFC North
    {"id": "BAL", "name": "Ravens", "city": "Baltimore", "division": "AFC North"},
    {"id": "CIN", "name": "Bengals", "city": "Cincinnati", "division": "AFC North"},
    {"id": "CLE", "name": "Browns", "city": "Cleveland", "division": "AFC North"},
    {"id": "PIT", "name": "Steelers", "city": "Pittsburgh", "division": "AFC North"},
    
    # AFC South
    {"id": "HOU", "name": "Texans", "city": "Houston", "division": "AFC South"},
    {"id": "IND", "name": "Colts", "city": "Indianapolis", "division": "AFC South"},
    {"id": "JAX", "name": "Jaguars", "city": "Jacksonville", "division": "AFC South"},
    {"id": "TEN", "name": "Titans", "city": "Tennessee", "division": "AFC South"},
    
    # AFC West
    {"id": "DEN", "name": "Broncos", "city": "Denver", "division": "AFC West"},
    {"id": "KC", "name": "Chiefs", "city": "Kansas City", "division": "AFC West"},
    {"id": "LV", "name": "Raiders", "city": "Las Vegas", "division": "AFC West"},
    {"id": "LAC", "name": "Chargers", "city": "Los Angeles", "division": "AFC West"},
    
    # NFC East
    {"id": "DAL", "name": "Cowboys", "city": "Dallas", "division": "NFC East"},
    {"id": "NYG", "name": "Giants", "city": "New York", "division": "NFC East"},
    {"id": "PHI", "name": "Eagles", "city": "Philadelphia", "division": "NFC East"},
    {"id": "WAS", "name": "Commanders", "city": "Washington", "division": "NFC East"},
    
    # NFC North
    {"id": "CHI", "name": "Bears", "city": "Chicago", "division": "NFC North"},
    {"id": "DET", "name": "Lions", "city": "Detroit", "division": "NFC North"},
    {"id": "GB", "name": "Packers", "city": "Green Bay", "division": "NFC North"},
    {"id": "MIN", "name": "Vikings", "city": "Minnesota", "division": "NFC North"},
    
    # NFC South
    {"id": "ATL", "name": "Falcons", "city": "Atlanta", "division": "NFC South"},
    {"id": "CAR", "name": "Panthers", "city": "Carolina", "division": "NFC South"},
    {"id": "NO", "name": "Saints", "city": "New Orleans", "division": "NFC South"},
    {"id": "TB", "name": "Buccaneers", "city": "Tampa Bay", "division": "NFC South"},
    
    # NFC West
    {"id": "ARI", "name": "Cardinals", "city": "Arizona", "division": "NFC West"},
    {"id": "LA", "name": "Rams", "city": "Los Angeles", "division": "NFC West"},
    {"id": "SF", "name": "49ers", "city": "San Francisco", "division": "NFC West"},
    {"id": "SEA", "name": "Seahawks", "city": "Seattle", "division": "NFC West"}
]

def convert_svg_to_png(svg_data, output_path):
    """
    SVG adat konvertálása PNG-re és mentése
    """
    try:
        if CAIROSVG_AVAILABLE:
            # CairoSVG használata
            if isinstance(svg_data, str):  # Ha fájl útvonal
                cairosvg.svg2png(url=svg_data, write_to=output_path)
            else:  # Ha bytes
                cairosvg.svg2png(bytestring=svg_data, write_to=output_path)
            return True
        else:
            # Ha nincs CairoSVG, akkor mentjük előbb az SVG-t, majd az inkscape-et használjuk
            if isinstance(svg_data, bytes):
                svg_path = output_path.replace('.png', '.svg')
                with open(svg_path, 'wb') as f:
                    f.write(svg_data)
                
                # Inkscape parancssoros hívása
                inkscape_cmd = f"inkscape --export-filename={output_path} {svg_path}"
                if os.system(inkscape_cmd) == 0 and os.path.exists(output_path):
                    return True
    except Exception as e:
        print(f"Hiba az SVG->PNG konvertálás során: {str(e)}")
    
    return False

def create_text_image(text, output_path, width=200, height=200):
    """
    Egyszerű szöveges kép létrehozása
    """
    try:
        # Létrehozunk egy üres képet
        img = PILImage.new('RGB', (width, height), color=(240, 240, 240))
        
        # Szöveg kirajzolásához betöltjük az ImageDraw modult
        from PIL import ImageDraw, ImageFont
        draw = ImageDraw.Draw(img)
        
        # Alapértelmezett betűtípus
        font = ImageFont.load_default()
        
        # Szöveg elhelyezése középre
        lines = text.split("\n")
        y_pos = (height - len(lines) * 20) // 2
        
        for line in lines:
            # Szöveghossz becslése (pontos mérés nem mindig elérhető)
            textwidth = font.getlength(line) if hasattr(font, 'getlength') else len(line) * 6
            x_pos = (width - textwidth) // 2
            
            # Szöveg kirajzolása
            draw.text((x_pos, y_pos), line, fill=(0, 0, 0))
            y_pos += 20
        
        # Kép mentése
        img.save(output_path)
        return True
    except Exception as e:
        print(f"Hiba a szöveges kép létrehozása közben: {str(e)}")
        return False

def get_logo_for_team(team_id, team_name, team_city, output_dir):
    """
    Adott csapathoz logó letöltése és visszaadása
    """
    # Meghatározzuk a kimeneti fájl nevét
    png_filename = os.path.join(output_dir, f"{team_id}_{team_name}.png")
    
    # NFL logó URL az adott csapathoz
    url = f"https://static.www.nfl.com/f_auto,q_85/league/api/clubs/logos/{team_id}"
    
    try:
        print(f"Attempting to download from URL: {url}")
        
        response = requests.get(url, stream=True)
        if response.status_code == 200:
            # Sikeres letöltés
            content_type = response.headers.get("Content-Type", "")
            
            if "svg" in content_type:
                # SVG fájl konvertálása PNG-re
                svg_data = response.content
                if convert_svg_to_png(svg_data, png_filename):
                    print(f"Logó letöltve és konvertálva: {team_city} {team_name} ({url})")
                    return png_filename
            else:
                # Közvetlenül PNG vagy más formátumú kép
                with open(png_filename, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        f.write(chunk)
                
                # Ellenőrizzük, hogy érvényes kép-e
                try:
                    with PILImage.open(png_filename) as img:
                        if img.width > 50 and img.height > 50:  # Minimális méret
                            print(f"Logó letöltve: {team_city} {team_name} ({url})")
                            return png_filename
                        else:
                            print(f"Túl kicsi a letöltött logó, próbáljuk alternatív megoldást...")
                except Exception as e:
                    print(f"A letöltött fájl nem érvényes kép: {str(e)}")
                    # Töröljük a hibás fájlt
                    if os.path.exists(png_filename):
                        os.remove(png_filename)
    except Exception as e:
        print(f"Hiba a logó letöltése közben ({url}): {str(e)}")
    
    # Ha nem sikerült a letöltés, próbáljunk alternatív forrásokat
    alternate_urls = [
        f"https://a.espncdn.com/i/teamlogos/nfl/500/scoreboard/{team_id.lower()}.png",
        f"https://a.espncdn.com/combiner/i?img=/i/teamlogos/nfl/500/{team_id.lower()}.png"
    ]
    
    for alt_url in alternate_urls:
        try:
            print(f"Attempting to download from alternate URL: {alt_url}")
            
            response = requests.get(alt_url, stream=True)
            if response.status_code == 200:
                # Közvetlenül PNG vagy más formátumú kép
                with open(png_filename, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        f.write(chunk)
                
                # Ellenőrizzük, hogy érvényes kép-e
                try:
                    with PILImage.open(png_filename) as img:
                        if img.width > 50 and img.height > 50:  # Minimális méret
                            print(f"Logó letöltve alternatív forrásból: {team_city} {team_name} ({alt_url})")
                            return png_filename
                except Exception:
                    # Töröljük a hibás fájlt
                    if os.path.exists(png_filename):
                        os.remove(png_filename)
        except Exception as e:
            print(f"Hiba az alternatív logó letöltése közben ({alt_url}): {str(e)}")
    
    # Ha minden forrás sikertelen volt, készítsünk egy szöveges helyettesítő képet
    print(f"Nem sikerült letölteni a logót: {team_city} {team_name}, helyettesítő kép készítése")
    create_text_image(f"{team_city}\n{team_name}", png_filename)
    
    return png_filename

def download_nfl_logos(output_dir="nfl_logos"):
    """
    NFL csapatok logóinak letöltése és helyettesítő képek készítése szükség esetén
    """
    # Létrehozzuk a kimeneti mappát
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    print(f"NFL logók letöltése a következő mappába: {output_dir}")
    
    # Eredmény követése
    download_results = {"successful": [], "failed": []}
    
    # Logók letöltése minden csapathoz
    for team in NFL_TEAMS:
        team_id = team["id"]
        team_name = team["name"]
        team_city = team["city"]
        
        # Logó letöltése
        logo_path = get_logo_for_team(team_id, team_name, team_city, output_dir)
        
        # Ellenőrizzük, hogy jó-e a logó
        success = False
        try:
            with PILImage.open(logo_path) as img:
                if img.width > 50 and img.height > 50:
                    success = True
        except Exception:
            success = False
        
        if success:
            download_results["successful"].append({
                "team_id": team_id,
                "name": f"{team_city} {team_name}",
                "division": team["division"],
                "logo_path": os.path.basename(logo_path)
            })
        else:
            download_results["failed"].append({
                "team_id": team_id,
                "name": f"{team_city} {team_name}",
                "division": team["division"]
            })
    
    # Mentsük el az eredményeket JSON-be
    results_file = os.path.join(output_dir, "download_results.json")
    with open(results_file, 'w', encoding='utf-8') as f:
        json.dump(download_results, f, indent=2, ensure_ascii=False)
    
    print(f"\nÖsszegzés:")
    print(f"Sikeresen letöltve: {len(download_results['successful'])} logó")
    print(f"Sikertelen: {len(download_results['failed'])} csapat")
    
    if download_results["failed"]:
        print(f"Sikertelen letöltések: {[t['name'] for t in download_results['failed']]}")
    
    return len(download_results["successful"])

def create_nfl_teams_word_document(logos_dir="nfl_logos", output_file="nfl_teams_logos.docx"):
    """
    Word dokumentum létrehozása NFL csapatok logóival
    """
    print(f"Word dokumentum létrehozása a következő mappából: {logos_dir}")
    
    # Ellenőrizzük, hogy létezik-e a mappa
    if not os.path.exists(logos_dir):
        print(f"A megadott mappa nem létezik: {logos_dir}")
        return False
    
    # Gyűjtsük össze az elérhető PNG/JPG képeket
    logo_files = []
    for root, _, files in os.walk(logos_dir):
        for file in files:
            if file.lower().endswith(('.png', '.jpg', '.jpeg')):
                logo_files.append(os.path.join(root, file))
    
    if not logo_files:
        print(f"Nem találtunk képfájlokat a következő mappában: {logos_dir}")
        return False
    
    print(f"Összesen {len(logo_files)} képet találtunk")
    
    # Hozzuk létre az új Word dokumentumot
    doc = Document()
    
    # Állítsuk be a dokumentum tulajdonságait
    doc.core_properties.title = "NFL Csapatok Logói"
    doc.core_properties.author = "NFL Logo Downloader"
    
    # Címlap
    title = doc.add_heading('NFL Csapatok', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Rövid magyarázat
    p = doc.add_paragraph('Az alábbi dokumentum tartalmazza az NFL csapatok logóit, konferenciáit és divízióit.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Gyűjtsük ki a csapatinformációkat a fájlnevekből és az NFL_TEAMS adatokból
    teams_info = {}
    
    # Először az NFL_TEAMS adatokat használjuk alapként
    for team in NFL_TEAMS:
        team_id = team["id"]
        teams_info[team_id] = {
            'name': team["name"],
            'city': team["city"],
            'division': team["division"],
            'logo_path': None  # Először nincs logó, majd a fájlnevek alapján beállítjuk
        }
    
    # Rendeljük hozzá a logó fájlokat a megfelelő csapatokhoz
    for logo_path in logo_files:
        file_name = os.path.basename(logo_path)
        # Várható formátum: ID_NÉV.png
        for team_id, team_info in teams_info.items():
            if file_name.startswith(f"{team_id}_"):
                teams_info[team_id]['logo_path'] = logo_path
                break
    
    # Ha még mindig vannak csapatok logó nélkül, készítsünk helyettesítő képeket
    for team_id, team_info in teams_info.items():
        if not team_info['logo_path']:
            team_name = team_info['name']
            team_city = team_info['city']
            
            print(f"Hiányzó logó pótlása: {team_city} {team_name}")
            placeholder_png = os.path.join(logos_dir, f"{team_id}_{team_name}.png")
            create_text_image(f"{team_city}\n{team_name}", placeholder_png)
            teams_info[team_id]['logo_path'] = placeholder_png
    
    # Rendezzük a csapatokat konferenciánként és divíziók szerint
    divisions = {
        "AFC East": [],
        "AFC North": [],
        "AFC South": [],
        "AFC West": [],
        "NFC East": [],
        "NFC North": [],
        "NFC South": [],
        "NFC West": []
    }
    
    # Csapatok rendezése divíziók szerint
    for team_id, team_info in teams_info.items():
        division = team_info['division']
        if division in divisions:
            divisions[division].append((team_id, team_info))
    
    # Dokumentum szélességének beállítása (hogy elférjen 4 oszlop)
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4 szélesség
    section.page_height = Cm(29.7)  # A4 magasság
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    
    # AFC konferencia
    doc.add_heading('AFC', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # AFC divíziók
    for division_name in ["AFC East", "AFC North", "AFC South", "AFC West"]:
        # Divízió címsor
        doc.add_heading(division_name, level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Táblázat a divízió csapataihoz
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Állítsuk be a táblázat oszlopainak szélességét
        for col_idx in range(4):
            for cell in table.columns[col_idx].cells:
                cell.width = Cm(4.5)
        
        # A divízió csapatai
        division_teams = divisions[division_name]
        division_teams.sort(key=lambda x: x[1]['city'])  # Rendezzük városnév szerint
        
        # Adjunk hozzá csapatokat a táblázathoz
        for col_idx, (team_id, team_info) in enumerate(division_teams):
            cell = table.cell(0, col_idx)
            team_name = team_info.get('name', '')
            team_city = team_info.get('city', '')
            logo_path = team_info.get('logo_path')
            
            if logo_path and os.path.exists(logo_path):
                # Hozzáadjuk a képet és a nevet
                p = cell.paragraphs[0]
                run = p.add_run()
                
                try:
                    # Beillesztjük a képet
                    run.add_picture(logo_path, width=Cm(2.5))
                except Exception as e:
                    print(f"Hiba a kép beillesztésekor ({team_name}): {str(e)}")
                    run.add_text(f"[{team_city} {team_name}]")
                
                # Adjuk hozzá a csapat nevét és városát
                name_text = f"{team_city} {team_name}"
                
                # Név hozzáadása a cellához
                name_paragraph = cell.add_paragraph(name_text)
                name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Beállítjuk a szöveg betűméretét 9-re
                for run in name_paragraph.runs:
                    run.font.size = Pt(9)
                
                # Csökkentsük a bekezdések közötti térközt
                name_paragraph.paragraph_format.space_after = Pt(2)
                
            else:
                # Ha nincs kép, csak a nevet és várost adjuk hozzá
                p = cell.paragraphs[0]
                p.text = f"{team_city} {team_name}"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Beállítjuk a szöveg betűméretét 9-re
                for run in p.runs:
                    run.font.size = Pt(9)
        
        # Téglalap kitöltése üres cellákkal, ha szükséges
        while len(division_teams) < 4:
            col_idx = len(division_teams)
            cell = table.cell(0, col_idx)
            cell.text = ""
            division_teams.append(None)  # Csak a hossz miatt
        
        # Üres sor a divíziók között
        doc.add_paragraph("")
    
    # Új oldal az NFC konferenciához
    doc.add_page_break()
    
    # NFC konferencia
    doc.add_heading('NFC', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # NFC divíziók
    for division_name in ["NFC East", "NFC North", "NFC South", "NFC West"]:
        # Divízió címsor
        doc.add_heading(division_name, level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Táblázat a divízió csapataihoz
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Állítsuk be a táblázat oszlopainak szélességét
        for col_idx in range(4):
            for cell in table.columns[col_idx].cells:
                cell.width = Cm(4.5)
        
        # A divízió csapatai
        division_teams = divisions[division_name]
        division_teams.sort(key=lambda x: x[1]['city'])  # Rendezzük városnév szerint
        
        # Adjunk hozzá csapatokat a táblázathoz
        for col_idx, (team_id, team_info) in enumerate(division_teams):
            cell = table.cell(0, col_idx)
            team_name = team_info.get('name', '')
            team_city = team_info.get('city', '')
            logo_path = team_info.get('logo_path')
            
            if logo_path and os.path.exists(logo_path):
                # Hozzáadjuk a képet és a nevet
                p = cell.paragraphs[0]
                run = p.add_run()
                
                try:
                    # Beillesztjük a képet
                    run.add_picture(logo_path, width=Cm(2.5))
                except Exception as e:
                    print(f"Hiba a kép beillesztésekor ({team_name}): {str(e)}")
                    run.add_text(f"[{team_city} {team_name}]")
                
                # Adjuk hozzá a csapat nevét és városát
                name_text = f"{team_city} {team_name}"
                
                # Név hozzáadása a cellához
                name_paragraph = cell.add_paragraph(name_text)
                name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Beállítjuk a szöveg betűméretét 9-re
                for run in name_paragraph.runs:
                    run.font.size = Pt(9)
                
                # Csökkentsük a bekezdések közötti térközt
                name_paragraph.paragraph_format.space_after = Pt(2)
                
            else:
                # Ha nincs kép, csak a nevet és várost adjuk hozzá
                p = cell.paragraphs[0]
                p.text = f"{team_city} {team_name}"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Beállítjuk a szöveg betűméretét 9-re
                for run in p.runs:
                    run.font.size = Pt(9)
        
        # Téglalap kitöltése üres cellákkal, ha szükséges
        while len(division_teams) < 4:
            col_idx = len(division_teams)
            cell = table.cell(0, col_idx)
            cell.text = ""
            division_teams.append(None)  # Csak a hossz miatt
        
        # Üres sor a divíziók között
        doc.add_paragraph("")
    
    # Mentés
    try:
        doc.save(output_file)
        print(f"A Word dokumentum sikeresen létrehozva: {output_file}")
        return True
    except Exception as e:
        print(f"Hiba a Word dokumentum mentésekor: {str(e)}")
        traceback.print_exc()
        return False

def main():
    # Beállítjuk a mappát, ahol a logókat tároljuk
    logos_dir = "nfl_logos"
    output_file = "nfl_teams_logos.docx"
    
    # Letöltjük a logókat
    print("Logók letöltése...")
    download_nfl_logos(logos_dir)
    
    # Word dokumentum létrehozása
    print("\nWord dokumentum készítése...")
    result = create_nfl_teams_word_document(logos_dir, output_file)
    
    if result:
        print(f"A dokumentum elkészült: {output_file}")
    else:
        print("A dokumentum létrehozása sikertelen.")

if __name__ == "__main__":
    main() 