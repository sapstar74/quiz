import os
import json
import sys
import requests
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import traceback
import urllib.parse
import io

# Próbáljuk importálni a CairoSVG-t az SVG konvertáláshoz
try:
    import cairosvg
    CAIROSVG_AVAILABLE = True
    print("CairoSVG sikeresen betöltve. Az SVG-k konvertálása lehetséges.")
except ImportError:
    CAIROSVG_AVAILABLE = False
    print("CairoSVG nem elérhető. Az SVG képek nem lesznek konvertálva PNG-re.")

# Az MLB csapatok definíciója a csapatkódokkal
MLB_TEAMS = [
    # American League East
    {"id": "110", "name": "Orioles", "city": "Baltimore"},
    {"id": "111", "name": "Red Sox", "city": "Boston"},
    {"id": "147", "name": "Yankees", "city": "New York"},
    {"id": "139", "name": "Rays", "city": "Tampa Bay"},
    {"id": "141", "name": "Blue Jays", "city": "Toronto"},
    
    # American League Central
    {"id": "145", "name": "White Sox", "city": "Chicago"},
    {"id": "114", "name": "Guardians", "city": "Cleveland"},
    {"id": "116", "name": "Tigers", "city": "Detroit"},
    {"id": "118", "name": "Royals", "city": "Kansas City"},
    {"id": "142", "name": "Twins", "city": "Minnesota"},
    
    # American League West
    {"id": "117", "name": "Astros", "city": "Houston"},
    {"id": "108", "name": "Angels", "city": "Los Angeles"},
    {"id": "133", "name": "Athletics", "city": "Oakland"},
    {"id": "136", "name": "Mariners", "city": "Seattle"},
    {"id": "140", "name": "Rangers", "city": "Texas"},
    
    # National League East
    {"id": "144", "name": "Braves", "city": "Atlanta"},
    {"id": "146", "name": "Marlins", "city": "Miami"},
    {"id": "121", "name": "Mets", "city": "New York"},
    {"id": "143", "name": "Phillies", "city": "Philadelphia"},
    {"id": "120", "name": "Nationals", "city": "Washington"},
    
    # National League Central
    {"id": "112", "name": "Cubs", "city": "Chicago"},
    {"id": "113", "name": "Reds", "city": "Cincinnati"},
    {"id": "158", "name": "Brewers", "city": "Milwaukee"},
    {"id": "134", "name": "Pirates", "city": "Pittsburgh"},
    {"id": "138", "name": "Cardinals", "city": "St. Louis"},
    
    # National League West
    {"id": "109", "name": "Diamondbacks", "city": "Arizona"},
    {"id": "115", "name": "Rockies", "city": "Colorado"},
    {"id": "119", "name": "Dodgers", "city": "Los Angeles"},
    {"id": "135", "name": "Padres", "city": "San Diego"},
    {"id": "137", "name": "Giants", "city": "San Francisco"},
]

def convert_svg_to_png(svg_data, output_path):
    """
    SVG adat konvertálása PNG-re és mentése
    """
    try:
        if CAIROSVG_AVAILABLE:
            # CairoSVG használata
            if isinstance(svg_data, str):  # Ha fájl útvonal
                cairosvg.svg2png(url=svg_data, write_to=output_path)
            else:  # Ha bytes
                cairosvg.svg2png(bytestring=svg_data, write_to=output_path)
            return True
        else:
            # Ha nincs CairoSVG, akkor mentjük előbb az SVG-t, majd az inkscape-et használjuk
            if isinstance(svg_data, bytes):
                svg_path = output_path.replace('.png', '.svg')
                with open(svg_path, 'wb') as f:
                    f.write(svg_data)
                
                # Inkscape parancssoros hívása
                inkscape_cmd = f"inkscape --export-filename={output_path} {svg_path}"
                if os.system(inkscape_cmd) == 0 and os.path.exists(output_path):
                    return True
    except Exception as e:
        print(f"Hiba az SVG->PNG konvertálás során: {str(e)}")
    
    return False

def create_text_image(text, output_path, width=200, height=200):
    """
    Egyszerű szöveges kép létrehozása
    """
    try:
        # Létrehozunk egy üres képet
        img = PILImage.new('RGB', (width, height), color=(240, 240, 240))
        
        # Szöveg kirajzolásához betöltjük az ImageDraw modult
        from PIL import ImageDraw, ImageFont
        draw = ImageDraw.Draw(img)
        
        # Alapértelmezett betűtípus
        font = ImageFont.load_default()
        
        # Szöveg elhelyezése középre
        lines = text.split("\n")
        y_pos = (height - len(lines) * 20) // 2
        
        for line in lines:
            # Szöveghossz becslése (pontos mérés nem mindig elérhető)
            textwidth = font.getlength(line) if hasattr(font, 'getlength') else len(line) * 6
            x_pos = (width - textwidth) // 2
            
            # Szöveg kirajzolása
            draw.text((x_pos, y_pos), line, fill=(0, 0, 0))
            y_pos += 20
        
        # Kép mentése
        img.save(output_path)
        return True
    except Exception as e:
        print(f"Hiba a szöveges kép létrehozása közben: {str(e)}")
        return False

def get_logo_for_team(team_id, team_name, team_city, output_dir):
    """
    Adott csapathoz logó letöltése és visszaadása
    """
    # Meghatározzuk a kimeneti fájl nevét
    png_filename = os.path.join(output_dir, f"{team_id}_{team_name}.png")
    
    # MLB logó URL az adott csapathoz
    url = f"https://www.mlbstatic.com/team-logos/team-cap-on-light/{team_id}.svg"
    
    try:
        print(f"Attempting to download from URL: {url}")
        
        response = requests.get(url, stream=True)
        if response.status_code == 200:
            # Sikeres letöltés
            content_type = response.headers.get("Content-Type", "")
            
            if "svg" in content_type:
                # SVG fájl konvertálása PNG-re
                svg_data = response.content
                if convert_svg_to_png(svg_data, png_filename):
                    print(f"Logó letöltve és konvertálva: {team_city} {team_name} ({url})")
                    return png_filename
            else:
                # Közvetlenül PNG vagy más formátumú kép
                with open(png_filename, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        f.write(chunk)
                
                # Ellenőrizzük, hogy érvényes kép-e
                try:
                    with PILImage.open(png_filename) as img:
                        if img.width > 50 and img.height > 50:  # Minimális méret
                            print(f"Logó letöltve: {team_city} {team_name} ({url})")
                            return png_filename
                        else:
                            print(f"Túl kicsi a letöltött logó, próbáljuk alternatív megoldást...")
                except Exception as e:
                    print(f"A letöltött fájl nem érvényes kép: {str(e)}")
                    # Töröljük a hibás fájlt
                    if os.path.exists(png_filename):
                        os.remove(png_filename)
    except Exception as e:
        print(f"Hiba a logó letöltése közben ({url}): {str(e)}")
    
    # Ha nem sikerült a letöltés, próbáljunk alternatív forrásokat
    alternate_urls = [
        f"https://www.mlbstatic.com/team-logos/{team_id}.svg",
        f"https://www.mlbstatic.com/team-logos/team-cap-on-dark/{team_id}.svg",
        f"https://content.mlb.com/images/logos/team-primary-on-light/{team_id}.svg",
        f"https://a.espncdn.com/i/teamlogos/mlb/500/scoreboard/{team_id.lower()}.png",
        f"https://a.espncdn.com/combiner/i?img=/i/teamlogos/mlb/500/{team_id.lower()}.png"
    ]
    
    for alt_url in alternate_urls:
        try:
            print(f"Attempting to download from alternate URL: {alt_url}")
            
            response = requests.get(alt_url, stream=True)
            if response.status_code == 200:
                content_type = response.headers.get("Content-Type", "")
                
                if "svg" in content_type:
                    # SVG fájl konvertálása PNG-re
                    svg_data = response.content
                    if convert_svg_to_png(svg_data, png_filename):
                        print(f"Logó letöltve és konvertálva alternatív forrásból: {team_city} {team_name} ({alt_url})")
                        return png_filename
                else:
                    # Közvetlenül PNG vagy más formátumú kép
                    with open(png_filename, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    
                    # Ellenőrizzük, hogy érvényes kép-e
                    try:
                        with PILImage.open(png_filename) as img:
                            if img.width > 50 and img.height > 50:  # Minimális méret
                                print(f"Logó letöltve alternatív forrásból: {team_city} {team_name} ({alt_url})")
                                return png_filename
                    except Exception:
                        # Töröljük a hibás fájlt
                        if os.path.exists(png_filename):
                            os.remove(png_filename)
        except Exception as e:
            print(f"Hiba az alternatív logó letöltése közben ({alt_url}): {str(e)}")
    
    # Ha minden forrás sikertelen volt, készítsünk egy szöveges helyettesítő képet
    print(f"Nem sikerült letölteni a logót: {team_city} {team_name}, helyettesítő kép készítése")
    create_text_image(f"{team_city}\n{team_name}", png_filename)
    
    return png_filename

def download_mlb_logos(output_dir="mlb_logos_flat"):
    """
    MLB csapatok logóinak letöltése és helyettesítő képek készítése szükség esetén
    """
    # Létrehozzuk a kimeneti mappát
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    print(f"MLB logók letöltése a következő mappába: {output_dir}")
    
    # Eredmény követése
    download_results = {"successful": [], "failed": []}
    
    # Logók letöltése minden csapathoz
    for team in MLB_TEAMS:
        team_id = team["id"]
        team_name = team["name"]
        team_city = team["city"]
        
        # Logó letöltése
        logo_path = get_logo_for_team(team_id, team_name, team_city, output_dir)
        
        # Ellenőrizzük, hogy jó-e a logó
        success = False
        try:
            with PILImage.open(logo_path) as img:
                if img.width > 50 and img.height > 50:
                    success = True
        except Exception:
            success = False
        
        if success:
            download_results["successful"].append({
                "team_id": team_id,
                "name": f"{team_city} {team_name}",
                "logo_path": os.path.basename(logo_path)
            })
        else:
            download_results["failed"].append({
                "team_id": team_id,
                "name": f"{team_city} {team_name}"
            })
    
    # Mentsük el az eredményeket JSON-be
    results_file = os.path.join(output_dir, "download_results.json")
    with open(results_file, 'w', encoding='utf-8') as f:
        json.dump(download_results, f, indent=2, ensure_ascii=False)
    
    print(f"\nÖsszegzés:")
    print(f"Sikeresen letöltve: {len(download_results['successful'])} logó")
    print(f"Sikertelen: {len(download_results['failed'])} csapat")
    
    if download_results["failed"]:
        print(f"Sikertelen letöltések: {[t['name'] for t in download_results['failed']]}")
    
    return len(download_results["successful"])

def create_mlb_teams_flat_document(logos_dir="mlb_logos_flat", output_file="mlb_teams_flat.docx"):
    """
    Word dokumentum létrehozása MLB csapatok logóival egy 8x4 táblázatban
    """
    print(f"Word dokumentum létrehozása a következő mappából: {logos_dir}")
    
    # Ellenőrizzük, hogy létezik-e a mappa
    if not os.path.exists(logos_dir):
        print(f"A megadott mappa nem létezik: {logos_dir}")
        return False
    
    # Gyűjtsük össze az elérhető PNG/JPG képeket
    logo_files = []
    for root, _, files in os.walk(logos_dir):
        for file in files:
            if file.lower().endswith(('.png', '.jpg', '.jpeg')):
                logo_files.append(os.path.join(root, file))
    
    if not logo_files:
        print(f"Nem találtunk képfájlokat a következő mappában: {logos_dir}")
        return False
    
    print(f"Összesen {len(logo_files)} képet találtunk")
    
    # Hozzuk létre az új Word dokumentumot
    doc = Document()
    
    # Állítsuk be a dokumentum tulajdonságait
    doc.core_properties.title = "MLB Csapatok Logói"
    doc.core_properties.author = "MLB Logo Downloader"
    
    # Címlap (egy sorban, egy kisebb címmel)
    title = doc.add_heading('MLB Csapatok', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dokumentum szélességének beállítása (hogy elférjen 4 oszlop)
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4 szélesség
    section.page_height = Cm(29.7)  # A4 magasság
    # Minimális margók, hogy maximalizáljuk a használható területet
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    
    # Gyűjtsük ki a csapatinformációkat a fájlnevekből és az MLB_TEAMS adatokból
    teams_info = {}
    
    # Először az MLB_TEAMS adatokat használjuk alapként
    for team in MLB_TEAMS:
        team_id = team["id"]
        teams_info[team_id] = {
            'name': team["name"],
            'city': team["city"],
            'logo_path': None  # Először nincs logó, majd a fájlnevek alapján beállítjuk
        }
    
    # Rendeljük hozzá a logó fájlokat a megfelelő csapatokhoz
    for logo_path in logo_files:
        file_name = os.path.basename(logo_path)
        # Várható formátum: ID_NÉV.png
        for team_id, team_info in teams_info.items():
            if file_name.startswith(f"{team_id}_"):
                teams_info[team_id]['logo_path'] = logo_path
                break
    
    # Ha még mindig vannak csapatok logó nélkül, készítsünk helyettesítő képeket
    for team_id, team_info in teams_info.items():
        if not team_info['logo_path']:
            team_name = team_info['name']
            team_city = team_info['city']
            
            print(f"Hiányzó logó pótlása: {team_city} {team_name}")
            placeholder_png = os.path.join(logos_dir, f"{team_id}_{team_name}.png")
            create_text_image(f"{team_city}\n{team_name}", placeholder_png)
            teams_info[team_id]['logo_path'] = placeholder_png
    
    # Táblázat létrehozása 8 sorral és 4 oszloppal
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    
    # Állítsuk be a táblázat teljes szélességét a lap szélességéhez
    table.autofit = False
    table.width = Cm(20)  # Az A4 szélesség mínusz a margók
    
    # Állítsuk be a táblázat oszlopainak szélességét egyenlő méretűre
    column_width = Cm(20 / 4)  # Egyenletesen osszuk el a szélességet
    for col_idx in range(4):
        for cell in table.columns[col_idx].cells:
            cell.width = column_width
    
    # A csapatokat alfabetikusan rendezzük a városnév alapján
    sorted_teams = sorted(teams_info.items(), key=lambda x: x[1]['city'])
    
    # Adjunk hozzá csapatokat a táblázathoz
    for idx, (team_id, team_info) in enumerate(sorted_teams):
        # Számoljuk ki a sor és oszlop pozíciót
        row_idx = idx // 4
        col_idx = idx % 4
        
        # Ha már nincs több sor a táblázatban, lépjünk ki
        if row_idx >= 8:
            print(f"Figyelem: Nem fér el minden csapat a táblázatban (8x4). A(z) {team_info['city']} {team_info['name']} kimarad.")
            continue
            
        cell = table.cell(row_idx, col_idx)
        team_name = team_info.get('name', '')
        team_city = team_info.get('city', '')
        logo_path = team_info.get('logo_path')
        
        # Kompaktabb cellákat készítünk
        cell.vertical_alignment = 1  # Középre igazítás vertikálisan
        
        if logo_path and os.path.exists(logo_path):
            # Hozzáadjuk a képet és a nevet
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(1)
            p.space_after = Pt(1)
            run = p.add_run()
            
            try:
                # Kisebb képméret a jobb elhelyezkedésért
                run.add_picture(logo_path, width=Cm(2.0))
            except Exception as e:
                print(f"Hiba a kép beillesztésekor ({team_name}): {str(e)}")
                run.add_text(f"[{team_city} {team_name}]")
            
            # Adjuk hozzá a csapat nevét és városát
            name_text = f"{team_city} {team_name}"
            
            # Név hozzáadása a cellához
            name_paragraph = cell.add_paragraph(name_text)
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Csökkentjük a térközt a bekezdések között
            name_paragraph.space_before = Pt(0)
            name_paragraph.space_after = Pt(0)
            
            # Beállítjuk a szöveg betűméretét 8-ra (kisebb, hogy jobban elférjen)
            for run in name_paragraph.runs:
                run.font.size = Pt(8)
            
        else:
            # Ha nincs kép, csak a nevet és várost adjuk hozzá
            p = cell.paragraphs[0]
            p.text = f"{team_city} {team_name}"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Beállítjuk a szöveg betűméretét 8-ra
            for run in p.runs:
                run.font.size = Pt(8)
    
    # A táblázat sorainak méretét optimalizáljuk
    for row in table.rows:
        row.height_rule = 1  # WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(3.2)  # Minden sor azonos magasságú
    
    # Mentés
    try:
        doc.save(output_file)
        print(f"A Word dokumentum sikeresen létrehozva: {output_file}")
        return True
    except Exception as e:
        print(f"Hiba a Word dokumentum mentésekor: {str(e)}")
        traceback.print_exc()
        return False

def main():
    # Beállítjuk a mappát, ahol a logókat tároljuk
    logos_dir = "mlb_logos_flat"
    output_file = "mlb_teams_flat.docx"
    
    # Letöltjük a logókat
    print("Logók letöltése...")
    download_mlb_logos(logos_dir)
    
    # Word dokumentum létrehozása
    print("\nWord dokumentum készítése...")
    result = create_mlb_teams_flat_document(logos_dir, output_file)
    
    if result:
        print(f"A dokumentum elkészült: {output_file}")
    else:
        print("A dokumentum létrehozása sikertelen.")

if __name__ == "__main__":
    main() 