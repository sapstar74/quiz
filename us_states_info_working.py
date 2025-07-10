import os
import requests
import io
from PIL import Image, ImageDraw, ImageFont
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import json

# Create directory for state seals
os.makedirs('us_state_seals', exist_ok=True)

# Dictionary of US states with their capitals
US_STATES = [
    {"name": "Alabama", "abbreviation": "AL", "capital": "Montgomery"},
    {"name": "Alaska", "abbreviation": "AK", "capital": "Juneau"},
    {"name": "Arizona", "abbreviation": "AZ", "capital": "Phoenix"},
    {"name": "Arkansas", "abbreviation": "AR", "capital": "Little Rock"},
    {"name": "California", "abbreviation": "CA", "capital": "Sacramento"},
    {"name": "Colorado", "abbreviation": "CO", "capital": "Denver"},
    {"name": "Connecticut", "abbreviation": "CT", "capital": "Hartford"},
    {"name": "Delaware", "abbreviation": "DE", "capital": "Dover"},
    {"name": "Florida", "abbreviation": "FL", "capital": "Tallahassee"},
    {"name": "Georgia", "abbreviation": "GA", "capital": "Atlanta"},
    {"name": "Hawaii", "abbreviation": "HI", "capital": "Honolulu"},
    {"name": "Idaho", "abbreviation": "ID", "capital": "Boise"},
    {"name": "Illinois", "abbreviation": "IL", "capital": "Springfield"},
    {"name": "Indiana", "abbreviation": "IN", "capital": "Indianapolis"},
    {"name": "Iowa", "abbreviation": "IA", "capital": "Des Moines"},
    {"name": "Kansas", "abbreviation": "KS", "capital": "Topeka"},
    {"name": "Kentucky", "abbreviation": "KY", "capital": "Frankfort"},
    {"name": "Louisiana", "abbreviation": "LA", "capital": "Baton Rouge"},
    {"name": "Maine", "abbreviation": "ME", "capital": "Augusta"},
    {"name": "Maryland", "abbreviation": "MD", "capital": "Annapolis"},
    {"name": "Massachusetts", "abbreviation": "MA", "capital": "Boston"},
    {"name": "Michigan", "abbreviation": "MI", "capital": "Lansing"},
    {"name": "Minnesota", "abbreviation": "MN", "capital": "Saint Paul"},
    {"name": "Mississippi", "abbreviation": "MS", "capital": "Jackson"},
    {"name": "Missouri", "abbreviation": "MO", "capital": "Jefferson City"},
    {"name": "Montana", "abbreviation": "MT", "capital": "Helena"},
    {"name": "Nebraska", "abbreviation": "NE", "capital": "Lincoln"},
    {"name": "Nevada", "abbreviation": "NV", "capital": "Carson City"},
    {"name": "New Hampshire", "abbreviation": "NH", "capital": "Concord"},
    {"name": "New Jersey", "abbreviation": "NJ", "capital": "Trenton"},
    {"name": "New Mexico", "abbreviation": "NM", "capital": "Santa Fe"},
    {"name": "New York", "abbreviation": "NY", "capital": "Albany"},
    {"name": "North Carolina", "abbreviation": "NC", "capital": "Raleigh"},
    {"name": "North Dakota", "abbreviation": "ND", "capital": "Bismarck"},
    {"name": "Ohio", "abbreviation": "OH", "capital": "Columbus"},
    {"name": "Oklahoma", "abbreviation": "OK", "capital": "Oklahoma City"},
    {"name": "Oregon", "abbreviation": "OR", "capital": "Salem"},
    {"name": "Pennsylvania", "abbreviation": "PA", "capital": "Harrisburg"},
    {"name": "Rhode Island", "abbreviation": "RI", "capital": "Providence"},
    {"name": "South Carolina", "abbreviation": "SC", "capital": "Columbia"},
    {"name": "South Dakota", "abbreviation": "SD", "capital": "Pierre"},
    {"name": "Tennessee", "abbreviation": "TN", "capital": "Nashville"},
    {"name": "Texas", "abbreviation": "TX", "capital": "Austin"},
    {"name": "Utah", "abbreviation": "UT", "capital": "Salt Lake City"},
    {"name": "Vermont", "abbreviation": "VT", "capital": "Montpelier"},
    {"name": "Virginia", "abbreviation": "VA", "capital": "Richmond"},
    {"name": "Washington", "abbreviation": "WA", "capital": "Olympia"},
    {"name": "West Virginia", "abbreviation": "WV", "capital": "Charleston"},
    {"name": "Wisconsin", "abbreviation": "WI", "capital": "Madison"},
    {"name": "Wyoming", "abbreviation": "WY", "capital": "Cheyenne"},
]

# Wikipedia URLs for state seals (direct image URLs)
STATE_SEAL_URLS = {
    "AL": "https://upload.wikimedia.org/wikipedia/commons/thumb/f/f7/Seal_of_Alabama.svg/300px-Seal_of_Alabama.svg.png",
    "AK": "https://upload.wikimedia.org/wikipedia/commons/thumb/9/96/State_Seal_of_Alaska.svg/300px-State_Seal_of_Alaska.svg.png",
    "AZ": "https://upload.wikimedia.org/wikipedia/commons/thumb/7/7e/Arizona-StateSeal.svg/300px-Arizona-StateSeal.svg.png",
    "AR": "https://upload.wikimedia.org/wikipedia/commons/thumb/a/a4/Seal_of_Arkansas.svg/300px-Seal_of_Arkansas.svg.png",
    "CA": "https://upload.wikimedia.org/wikipedia/commons/thumb/0/0f/Seal_of_California.svg/300px-Seal_of_California.svg.png",
    "CO": "https://upload.wikimedia.org/wikipedia/commons/thumb/0/00/Seal_of_Colorado.svg/300px-Seal_of_Colorado.svg.png",
    "CT": "https://upload.wikimedia.org/wikipedia/commons/thumb/4/48/Seal_of_Connecticut.svg/300px-Seal_of_Connecticut.svg.png",
    "DE": "https://upload.wikimedia.org/wikipedia/commons/thumb/3/35/Seal_of_Delaware.svg/300px-Seal_of_Delaware.svg.png",
    "FL": "https://upload.wikimedia.org/wikipedia/commons/thumb/2/2b/Seal_of_Florida.svg/300px-Seal_of_Florida.svg.png",
    "GA": "https://upload.wikimedia.org/wikipedia/commons/thumb/7/79/Seal_of_Georgia.svg/300px-Seal_of_Georgia.svg.png",
    "HI": "https://upload.wikimedia.org/wikipedia/commons/thumb/c/ca/Seal_of_the_State_of_Hawaii.svg/300px-Seal_of_the_State_of_Hawaii.svg.png",
    "ID": "https://upload.wikimedia.org/wikipedia/commons/thumb/4/46/Seal_of_Idaho.svg/300px-Seal_of_Idaho.svg.png",
    "IL": "https://upload.wikimedia.org/wikipedia/commons/thumb/e/e7/Seal_of_Illinois.svg/300px-Seal_of_Illinois.svg.png",
    "IN": "https://upload.wikimedia.org/wikipedia/commons/thumb/c/c4/Indiana-StateSeal.svg/300px-Indiana-StateSeal.svg.png",
    "IA": "https://upload.wikimedia.org/wikipedia/commons/thumb/5/5a/Iowa-StateSeal.svg/300px-Iowa-StateSeal.svg.png",
    "KS": "https://upload.wikimedia.org/wikipedia/commons/thumb/4/45/Seal_of_Kansas.svg/300px-Seal_of_Kansas.svg.png",
    "KY": "https://upload.wikimedia.org/wikipedia/commons/thumb/3/35/Seal_of_Kentucky.svg/300px-Seal_of_Kentucky.svg.png",
    "LA": "https://upload.wikimedia.org/wikipedia/commons/thumb/2/2f/Seal_of_Louisiana.svg/300px-Seal_of_Louisiana.svg.png",
    "ME": "https://upload.wikimedia.org/wikipedia/commons/thumb/6/63/Seal_of_Maine.svg/300px-Seal_of_Maine.svg.png",
    "MD": "https://upload.wikimedia.org/wikipedia/commons/thumb/1/1a/Seal_of_Maryland_%28obverse%29.png/300px-Seal_of_Maryland_%28obverse%29.png",
    "MA": "https://upload.wikimedia.org/wikipedia/commons/thumb/8/82/Seal_of_Massachusetts.svg/300px-Seal_of_Massachusetts.svg.png",
    "MI": "https://upload.wikimedia.org/wikipedia/commons/thumb/3/3f/Seal_of_Michigan.svg/300px-Seal_of_Michigan.svg.png",
    "MN": "https://upload.wikimedia.org/wikipedia/commons/thumb/b/b9/Seal_of_Minnesota.svg/300px-Seal_of_Minnesota.svg.png",
    "MS": "https://upload.wikimedia.org/wikipedia/commons/thumb/c/cf/Seal_of_Mississippi_%282014%29.svg/300px-Seal_of_Mississippi_%282014%29.svg.png",
    "MO": "https://upload.wikimedia.org/wikipedia/commons/thumb/d/de/Seal_of_Missouri.svg/300px-Seal_of_Missouri.svg.png",
    "MT": "https://upload.wikimedia.org/wikipedia/commons/thumb/4/45/Montana_state_seal.png/300px-Montana_state_seal.png",
    "NE": "https://upload.wikimedia.org/wikipedia/commons/thumb/7/73/Seal_of_Nebraska.svg/300px-Seal_of_Nebraska.svg.png",
    "NV": "https://upload.wikimedia.org/wikipedia/commons/thumb/7/77/Nevada-StateSeal.svg/300px-Nevada-StateSeal.svg.png",
    "NH": "https://upload.wikimedia.org/wikipedia/commons/thumb/a/aa/Seal_of_New_Hampshire.svg/300px-Seal_of_New_Hampshire.svg.png",
    "NJ": "https://upload.wikimedia.org/wikipedia/commons/thumb/8/8e/Seal_of_New_Jersey.svg/300px-Seal_of_New_Jersey.svg.png",
    "NM": "https://upload.wikimedia.org/wikipedia/commons/thumb/4/47/Great_seal_of_the_state_of_New_Mexico.png/300px-Great_seal_of_the_state_of_New_Mexico.png",
    "NY": "https://upload.wikimedia.org/wikipedia/commons/thumb/c/ca/Seal_of_New_York.svg/300px-Seal_of_New_York.svg.png",
    "NC": "https://upload.wikimedia.org/wikipedia/commons/thumb/7/72/Seal_of_North_Carolina.svg/300px-Seal_of_North_Carolina.svg.png",
    "ND": "https://upload.wikimedia.org/wikipedia/commons/thumb/e/e7/NorthDakotaStateSeal.svg/300px-NorthDakotaStateSeal.svg.png",
    "OH": "https://upload.wikimedia.org/wikipedia/commons/thumb/f/f3/Seal_of_Ohio.svg/300px-Seal_of_Ohio.svg.png",
    "OK": "https://upload.wikimedia.org/wikipedia/commons/thumb/3/39/Seal_of_Oklahoma.svg/300px-Seal_of_Oklahoma.svg.png",
    "OR": "https://upload.wikimedia.org/wikipedia/commons/thumb/4/46/Seal_of_Oregon.svg/300px-Seal_of_Oregon.svg.png",
    "PA": "https://upload.wikimedia.org/wikipedia/commons/thumb/7/70/Seal_of_Pennsylvania.svg/300px-Seal_of_Pennsylvania.svg.png",
    "RI": "https://upload.wikimedia.org/wikipedia/commons/thumb/7/76/Seal_of_Rhode_Island.svg/300px-Seal_of_Rhode_Island.svg.png",
    "SC": "https://upload.wikimedia.org/wikipedia/commons/thumb/8/80/Seal_of_South_Carolina.svg/300px-Seal_of_South_Carolina.svg.png",
    "SD": "https://upload.wikimedia.org/wikipedia/commons/thumb/b/bb/SouthDakotaStateSeal.svg/300px-SouthDakotaStateSeal.svg.png",
    "TN": "https://upload.wikimedia.org/wikipedia/commons/thumb/3/3c/Seal_of_Tennessee.svg/300px-Seal_of_Tennessee.svg.png",
    "TX": "https://upload.wikimedia.org/wikipedia/commons/thumb/c/cb/Seal_of_Texas.svg/300px-Seal_of_Texas.svg.png",
    "UT": "https://upload.wikimedia.org/wikipedia/commons/thumb/a/a9/Seal_of_Utah.svg/300px-Seal_of_Utah.svg.png",
    "VT": "https://upload.wikimedia.org/wikipedia/commons/thumb/f/f9/Seal_of_Vermont.svg/300px-Seal_of_Vermont.svg.png",
    "VA": "https://upload.wikimedia.org/wikipedia/commons/thumb/5/5d/Seal_of_Virginia.svg/300px-Seal_of_Virginia.svg.png",
    "WA": "https://upload.wikimedia.org/wikipedia/commons/thumb/3/3d/Seal_of_Washington.svg/300px-Seal_of_Washington.svg.png",
    "WV": "https://upload.wikimedia.org/wikipedia/commons/thumb/1/1c/Seal_of_West_Virginia.svg/300px-Seal_of_West_Virginia.svg.png",
    "WI": "https://upload.wikimedia.org/wikipedia/commons/thumb/7/7f/Seal_of_Wisconsin.svg/300px-Seal_of_Wisconsin.svg.png",
    "WY": "https://upload.wikimedia.org/wikipedia/commons/thumb/e/e4/Seal_of_Wyoming.svg/300px-Seal_of_Wyoming.svg.png",
}

def create_text_image(text, size=(200, 200), bg_color=(240, 240, 240), text_color=(0, 0, 0)):
    """Create a text image as a placeholder for missing state seals."""
    img = Image.new('RGB', size, color=bg_color)
    draw = ImageDraw.Draw(img)
    
    # Use default font
    font = ImageFont.load_default()
    
    # Draw text centered (without using getsize which is deprecated)
    text_lines = text.split('\n')
    y_position = size[1] // 2 - 10 * len(text_lines)
    
    for line in text_lines:
        # Approximate center position
        x_position = size[0] // 2 - 5 * len(line)  
        draw.text((x_position, y_position), line, font=font, fill=text_color)
        y_position += 20  # Move down for next line
    
    return img

def download_state_seal(state_name, state_abbr):
    """Download a state seal image from pre-defined URL."""
    img_path = os.path.join('us_state_seals', f"{state_abbr}.png")
    
    # If image already exists, return its path
    if os.path.exists(img_path):
        return img_path
    
    # Get URL from pre-defined dict
    if state_abbr in STATE_SEAL_URLS:
        try:
            url = STATE_SEAL_URLS[state_abbr]
            print(f"Downloading seal for {state_name} from {url}")
            response = requests.get(url, timeout=10)
            
            if response.status_code == 200:
                # Save the image directly from response
                with open(img_path, 'wb') as f:
                    f.write(response.content)
                print(f"Successfully downloaded seal for {state_name}")
                return img_path
        except Exception as e:
            print(f"Error downloading seal for {state_name}: {e}")
    
    # If download fails, create a placeholder image
    try:
        placeholder = create_text_image(f"Seal of\n{state_name}")
        placeholder.save(img_path)
        print(f"Created placeholder seal for {state_name}")
        return img_path
    except Exception as e:
        print(f"Error creating placeholder for {state_name}: {e}")
        return None

def download_all_state_seals():
    """Download all state seals and return results."""
    results = {"success": [], "failed": []}
    
    for state in US_STATES:
        state_name = state["name"]
        state_abbr = state["abbreviation"]
        
        img_path = download_state_seal(state_name, state_abbr)
        
        if img_path and os.path.exists(img_path):
            results["success"].append({
                "state": state_name,
                "abbr": state_abbr,
                "path": img_path
            })
        else:
            results["failed"].append({
                "state": state_name,
                "abbr": state_abbr
            })
        
        # Add a small delay to avoid overwhelming servers
        time.sleep(0.2)
    
    # Save results to a JSON file for reference
    with open('us_state_seals/download_results.json', 'w') as f:
        json.dump(results, f, indent=2)
    
    return results

def create_us_states_doc():
    """Create a Word document with all US states, their seals, and capitals."""
    # First download all state seals
    print("Downloading state seals...")
    results = download_all_state_seals()
    print(f"Downloaded {len(results['success'])} state seals successfully")
    if results['failed']:
        print(f"Failed to download {len(results['failed'])} state seals")
    
    # Create a Word document
    doc = docx.Document()
    
    # Set document properties
    doc.core_properties.title = "United States - States and Capitals"
    
    # Set document margins (smaller margins for better fit)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
    
    # Add a title
    title = doc.add_heading("United States - States and Capitals", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create a table with 5 states per row
    # Total rows needed: ceil(50/5) = 10 rows
    table = doc.add_table(rows=10, cols=5)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set table width to fill page
    table.width = Cm(19)
    
    # Set equal width for each column
    for cell in table.columns:
        cell.width = Cm(3.8)
    
    # Add states to the table
    for i, state in enumerate(US_STATES):
        # Calculate row and column position
        row = i // 5
        col = i % 5
        
        # Get cell
        cell = table.cell(row, col)
        
        # Add state name as cell heading
        state_para = cell.paragraphs[0]
        state_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        state_run = state_para.add_run(state["name"])
        state_run.bold = True
        state_run.font.size = Pt(10)
        
        # Add state seal image
        img_path = os.path.join('us_state_seals', f"{state['abbreviation']}.png")
        if os.path.exists(img_path):
            try:
                # Add image paragraph
                img_para = cell.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(img_path, width=Cm(3.2))
            except Exception as e:
                print(f"Error adding image for {state['name']}: {e}")
                # Add a placeholder text instead
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("[Seal Image]")
        
        # Add capital city
        capital_para = cell.add_paragraph()
        capital_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        capital_run = capital_para.add_run(f"Capital: {state['capital']}")
        capital_run.font.size = Pt(9)
        
        # Set row height to accommodate content
        table.rows[row].height = Cm(4.5)
        cell.vertical_alignment = 1  # centered
    
    # Save the document
    output_file = "us_states_capitals.docx"
    doc.save(output_file)
    print(f"Document created successfully: {output_file}")
    return output_file

if __name__ == "__main__":
    create_us_states_doc() 