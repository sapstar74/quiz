import os
import json
import sys
import requests
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import traceback
import urllib.parse

# Próbáljuk importálni a CairoSVG-t az SVG konvertáláshoz
try:
    import cairosvg
    CAIROSVG_AVAILABLE = True
    print("CairoSVG sikeresen betöltve. Az SVG-k konvertálása lehetséges.")
except ImportError:
    CAIROSVG_AVAILABLE = False
    print("CairoSVG nem elérhető. Az SVG képek nem lesznek konvertálva PNG-re.")

# Az NBA csapatok definíciója a csapatkódokkal
NBA_TEAMS = [
    {"id": 1610612738, "abbreviation": "BOS", "name": "Celtics", "city": "Boston"},
    {"id": 1610612751, "abbreviation": "BKN", "name": "Nets", "city": "Brooklyn"},
    {"id": 1610612752, "abbreviation": "NYK", "name": "Knicks", "city": "New York"},
    {"id": 1610612755, "abbreviation": "PHI", "name": "76ers", "city": "Philadelphia"},
    {"id": 1610612761, "abbreviation": "TOR", "name": "Raptors", "city": "Toronto"},
    {"id": 1610612741, "abbreviation": "CHI", "name": "Bulls", "city": "Chicago"},
    {"id": 1610612739, "abbreviation": "CLE", "name": "Cavaliers", "city": "Cleveland"},
    {"id": 1610612765, "abbreviation": "DET", "name": "Pistons", "city": "Detroit"},
    {"id": 1610612754, "abbreviation": "IND", "name": "Pacers", "city": "Indiana"},
    {"id": 1610612749, "abbreviation": "MIL", "name": "Bucks", "city": "Milwaukee"},
    {"id": 1610612737, "abbreviation": "ATL", "name": "Hawks", "city": "Atlanta"},
    {"id": 1610612766, "abbreviation": "CHA", "name": "Hornets", "city": "Charlotte"},
    {"id": 1610612748, "abbreviation": "MIA", "name": "Heat", "city": "Miami"},
    {"id": 1610612753, "abbreviation": "ORL", "name": "Magic", "city": "Orlando"},
    {"id": 1610612764, "abbreviation": "WAS", "name": "Wizards", "city": "Washington"},
    {"id": 1610612743, "abbreviation": "DEN", "name": "Nuggets", "city": "Denver"},
    {"id": 1610612750, "abbreviation": "MIN", "name": "Timberwolves", "city": "Minnesota"},
    {"id": 1610612760, "abbreviation": "OKC", "name": "Thunder", "city": "Oklahoma City"},
    {"id": 1610612757, "abbreviation": "POR", "name": "Trail Blazers", "city": "Portland"},
    {"id": 1610612762, "abbreviation": "UTA", "name": "Jazz", "city": "Utah"},
    {"id": 1610612744, "abbreviation": "GSW", "name": "Warriors", "city": "Golden State"},
    {"id": 1610612746, "abbreviation": "LAC", "name": "Clippers", "city": "Los Angeles"},
    {"id": 1610612747, "abbreviation": "LAL", "name": "Lakers", "city": "Los Angeles"},
    {"id": 1610612756, "abbreviation": "PHX", "name": "Suns", "city": "Phoenix"},
    {"id": 1610612758, "abbreviation": "SAC", "name": "Kings", "city": "Sacramento"},
    {"id": 1610612742, "abbreviation": "DAL", "name": "Mavericks", "city": "Dallas"},
    {"id": 1610612745, "abbreviation": "HOU", "name": "Rockets", "city": "Houston"},
    {"id": 1610612763, "abbreviation": "MEM", "name": "Grizzlies", "city": "Memphis"},
    {"id": 1610612740, "abbreviation": "NOP", "name": "Pelicans", "city": "New Orleans"},
    {"id": 1610612759, "abbreviation": "SAS", "name": "Spurs", "city": "San Antonio"}
]

def convert_svg_to_png(svg_path, png_path):
    """
    SVG fájl konvertálása PNG-re (ha lehetséges)
    """
    try:
        if CAIROSVG_AVAILABLE:
            # CairoSVG használata
            cairosvg.svg2png(url=svg_path, write_to=png_path)
            return True
        else:
            # Inkscape parancssoros hívása alternatívaként
            inkscape_cmd = f"inkscape --export-filename={png_path} {svg_path}"
            if os.system(inkscape_cmd) == 0 and os.path.exists(png_path):
                return True
    except Exception as e:
        print(f"Hiba az SVG->PNG konvertálás során: {str(e)}")
    
    return False

def create_text_image(text, output_path, width=200, height=200):
    """
    Egyszerű szöveges kép létrehozása
    """
    try:
        # Létrehozunk egy üres képet
        img = PILImage.new('RGB', (width, height), color=(240, 240, 240))
        
        # Szöveg kirajzolásához betöltjük az ImageDraw modult
        from PIL import ImageDraw, ImageFont
        draw = ImageDraw.Draw(img)
        
        # Alapértelmezett betűtípus
        font = ImageFont.load_default()
        
        # Szöveg elhelyezése középre
        lines = text.split("\n")
        y_pos = (height - len(lines) * 20) // 2
        
        for line in lines:
            # Szöveghossz becslése (pontos mérés nem mindig elérhető)
            textwidth = font.getlength(line) if hasattr(font, 'getlength') else len(line) * 6
            x_pos = (width - textwidth) // 2
            
            # Szöveg kirajzolása
            draw.text((x_pos, y_pos), line, fill=(0, 0, 0))
            y_pos += 20
        
        # Kép mentése
        img.save(output_path)
        return True
    except Exception as e:
        print(f"Hiba a szöveges kép létrehozása közben: {str(e)}")
        return False

def download_nba_logos(output_dir="nba_logos"):
    """
    NBA csapatok logóinak letöltése a hivatalos CDN-ről
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    base_url = "https://cdn.nba.com/logos/nba/{team_id}/global/L/logo.svg"
    # Alternatív URL formátumok
    alt_url_formats = [
        "https://cdn.nba.com/logos/nba/{team_id}/primary/L/logo.svg",
        "https://cdn.nba.com/logos/nba/{team_id}/primary/D/logo.svg",
        "https://cdn.nba.com/logos/nba/{team_id}/global/D/logo.svg"
    ]
    
    # Problémás csapatok külön URL-jei
    special_urls = {
        1610612744: "https://upload.wikimedia.org/wikipedia/en/0/01/Golden_State_Warriors_logo.svg",  # Warriors
        1610612747: "https://upload.wikimedia.org/wikipedia/commons/3/3c/Los_Angeles_Lakers_logo.svg",     # Lakers
        1610612746: "https://upload.wikimedia.org/wikipedia/en/b/bb/Los_Angeles_Clippers_%282015%29.svg",  # Clippers
        1610612756: "https://upload.wikimedia.org/wikipedia/en/d/dc/Phoenix_Suns_logo.svg",            # Suns
        1610612742: "https://upload.wikimedia.org/wikipedia/en/9/97/Dallas_Mavericks_logo.svg",         # Mavericks
        1610612745: "https://upload.wikimedia.org/wikipedia/en/2/28/Houston_Rockets.svg",        # Rockets
        1610612763: "https://upload.wikimedia.org/wikipedia/en/f/f1/Memphis_Grizzlies.svg",      # Grizzlies
        1610612740: "https://upload.wikimedia.org/wikipedia/en/0/0d/New_Orleans_Pelicans_logo.svg", # Pelicans
        1610612759: "https://upload.wikimedia.org/wikipedia/en/a/a2/San_Antonio_Spurs.svg"       # Spurs
    }
    
    download_results = {
        "successful_downloads": {},
        "failed_downloads": {}
    }
    
    print(f"NBA logók letöltése a következő mappába: {output_dir}")
    
    for team in NBA_TEAMS:
        team_id = team["id"]
        team_abbreviation = team["abbreviation"]
        team_name = team["name"]
        team_city = team["city"]
        
        # Fájlnév meghatározása
        local_filename = os.path.join(output_dir, f"{team_id}_{team_abbreviation}_{team_name}.svg")
        png_filename = local_filename.replace('.svg', '.png')
        
        # Logó letöltése - először próbáljuk a speciális URL-t a problémás csapatoknál
        success = False
        
        if team_id in special_urls:
            print(f"Különleges URL használata: {team_city} {team_name}")
            special_url = special_urls[team_id]
            try:
                # Speciális esetben egyből PNG-t töltünk le
                special_filename = os.path.join(output_dir, f"{team_id}_{team_abbreviation}_{team_name}.png")
                response = requests.get(special_url, stream=True)
                
                if response.status_code == 200:
                    with open(special_filename, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    
                    print(f"Speciális URL-ről letöltve: {team_city} {team_name} - {special_filename}")
                    success = True
                    
                    # Sikeresen letöltött logók közé adjuk
                    download_results["successful_downloads"][str(team_id)] = {
                        "team": f"{team_city} {team_name}",
                        "abbreviation": team_abbreviation,
                        "path": os.path.basename(special_filename)
                    }
            except Exception as e:
                print(f"Hiba a speciális URL-ről való letöltéskor: {team_city} {team_name} - {str(e)}")
        
        # Ha nem sikerült a speciális URL-ről, próbáljuk az alap és alternatív URL-eket
        if not success:
            # Próbáljuk az alap URL-t
            logo_url = base_url.format(team_id=team_id)
            try:
                response = requests.get(logo_url, stream=True)
                
                if response.status_code == 200:
                    with open(local_filename, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    
                    print(f"Letöltve: {team_city} {team_name} - {local_filename}")
                    
                    # SVG konvertálása PNG-re
                    if convert_svg_to_png(local_filename, png_filename):
                        print(f"  SVG konvertálva PNG-re: {png_filename}")
                    
                    # Sikeresen letöltött logók közé adjuk
                    download_results["successful_downloads"][str(team_id)] = {
                        "team": f"{team_city} {team_name}",
                        "abbreviation": team_abbreviation,
                        "path": os.path.basename(local_filename)
                    }
                    success = True
                else:
                    print(f"Az alap URL sikertelen ({response.status_code}), alternatív URL-ek próbálása: {team_city} {team_name}")
                    # Próbáljuk az alternatív URL-eket
                    for alt_url_format in alt_url_formats:
                        alt_url = alt_url_format.format(team_id=team_id)
                        try:
                            alt_response = requests.get(alt_url, stream=True)
                            if alt_response.status_code == 200:
                                with open(local_filename, 'wb') as f:
                                    for chunk in alt_response.iter_content(chunk_size=8192):
                                        f.write(chunk)
                                
                                print(f"Alternatív URL-ről letöltve: {team_city} {team_name} - {local_filename}")
                                
                                # SVG konvertálása PNG-re
                                if convert_svg_to_png(local_filename, png_filename):
                                    print(f"  SVG konvertálva PNG-re: {png_filename}")
                                
                                # Sikeresen letöltött logók közé adjuk
                                download_results["successful_downloads"][str(team_id)] = {
                                    "team": f"{team_city} {team_name}",
                                    "abbreviation": team_abbreviation,
                                    "path": os.path.basename(local_filename)
                                }
                                success = True
                                break
                        except Exception as e:
                            print(f"Hiba az alternatív URL próbálásakor: {alt_url} - {str(e)}")
            except Exception as e:
                print(f"Hiba a letöltéskor: {team_city} {team_name} - {str(e)}")
        
        # Ha még mindig nem sikerült, készítsünk helyette egy szöveges képet
        if not success:
            print(f"Nem sikerült letölteni {team_city} {team_name} logóját, helyettesítő kép készítése")
            download_results["failed_downloads"][str(team_id)] = {
                "team": f"{team_city} {team_name}",
                "abbreviation": team_abbreviation,
                "error": "Minden URL próba sikertelen"
            }
            
            placeholder_png = os.path.join(output_dir, f"{team_id}_{team_abbreviation}_{team_name}.png")
            create_text_image(f"{team_city}\n{team_name}", placeholder_png)
            print(f"  Helyettesítő kép készült: {placeholder_png}")
    
    # Mentsük el az eredményeket JSON-be
    results_file = os.path.join(output_dir, "download_results.json")
    with open(results_file, 'w', encoding='utf-8') as f:
        json.dump(download_results, f, indent=2, ensure_ascii=False)
    
    print(f"Letöltési eredmények mentve: {results_file}")
    print(f"Sikeresen letöltve: {len(download_results['successful_downloads'])} logó")
    print(f"Sikertelenül: {len(download_results['failed_downloads'])} logó")
    
    return download_results

def create_nba_teams_word_document(logos_dir="nba_logos", output_file="nba_teams_logos.docx"):
    """
    Word dokumentum létrehozása NBA csapatok logóival
    """
    print(f"Word dokumentum létrehozása a következő mappából: {logos_dir}")
    
    # Ellenőrizzük, hogy létezik-e a mappa
    if not os.path.exists(logos_dir):
        print(f"A megadott mappa nem létezik: {logos_dir}")
        return False
    
    # Nézzük meg, van-e letöltött információnk a logókról
    results_file = os.path.join(logos_dir, "download_results.json")
    results_data = {}
    
    if os.path.exists(results_file):
        try:
            with open(results_file, 'r', encoding='utf-8') as f:
                results_data = json.load(f)
            print(f"Sikeresen betöltöttük a letöltési eredményeket: {len(results_data.get('successful_downloads', {}))} csapat")
        except Exception as e:
            print(f"Hiba a letöltési eredmények betöltésekor: {str(e)}")
    
    # SVG képek konvertálása PNG-re (Word nem kezeli az SVG-ket)
    svg_files = []
    for root, _, files in os.walk(logos_dir):
        for file in files:
            if file.lower().endswith('.svg'):
                svg_path = os.path.join(root, file)
                png_path = svg_path.replace('.svg', '.png')
                
                if not os.path.exists(png_path):
                    print(f"SVG konvertálása PNG-re: {file}")
                    if not convert_svg_to_png(svg_path, png_path):
                        # Ha a konvertálás sikertelen, készítsünk egy szöveges képet helyette
                        parts = file.split('_')
                        if len(parts) >= 3:
                            team_name = parts[2].replace('.svg', '')
                            team_abbr = parts[1]
                            team_info = next((team for team in NBA_TEAMS if team["abbreviation"] == team_abbr), None)
                            
                            if team_info:
                                text = f"{team_info['city']}\n{team_info['name']}"
                            else:
                                text = team_name
                                
                            print(f"  Helyettesítő kép készítése: {text}")
                            create_text_image(text, png_path)
    
    # Gyűjtsük össze az elérhető PNG/JPG képeket
    logo_files = []
    for root, _, files in os.walk(logos_dir):
        for file in files:
            if file.lower().endswith(('.png', '.jpg', '.jpeg')):
                logo_files.append(os.path.join(root, file))
    
    # Ellenőrizzük az utolsó sor teamjeit, és szükség esetén készítsünk helyettesítő képeket
    last_row_abbrs = ["DAL", "HOU", "MEM", "NOP", "SAS"]
    for team in NBA_TEAMS:
        if team["abbreviation"] in last_row_abbrs:
            team_id = team["id"]
            team_name = team["name"]
            team_abbreviation = team["abbreviation"]
            team_city = team["city"]
            
            # Ellenőrizzük, van-e megfelelő kép
            has_valid_logo = False
            for logo_path in logo_files:
                if f"{team_id}_{team_abbreviation}_" in os.path.basename(logo_path):
                    # Ellenőrizzük a kép minőségét/méretét
                    try:
                        with PILImage.open(logo_path) as img:
                            if img.width > 50 and img.height > 50:  # Minimális méret
                                has_valid_logo = True
                                break
                    except Exception as e:
                        print(f"Hiba a kép ellenőrzésekor ({team_name}): {str(e)}")
            
            # Ha nincs megfelelő kép, készítsünk helyettesítőt
            if not has_valid_logo:
                print(f"Helyettesítő kép készítése: {team_city} {team_name}")
                placeholder_png = os.path.join(logos_dir, f"{team_id}_{team_abbreviation}_{team_name}.png")
                create_text_image(f"{team_city}\n{team_name}", placeholder_png)
                if placeholder_png not in logo_files:
                    logo_files.append(placeholder_png)
    
    if not logo_files:
        print(f"Nem találtunk képfájlokat a következő mappában: {logos_dir}")
        return False
    
    print(f"Összesen {len(logo_files)} képet találtunk")
    
    # Hozzuk létre az új Word dokumentumot
    doc = Document()
    
    # Állítsuk be a dokumentum tulajdonságait
    doc.core_properties.title = "NBA Csapatok Logói"
    doc.core_properties.author = "NBA Logo Downloader"
    
    # Címlap
    title = doc.add_heading('NBA Csapatok', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Rövid magyarázat
    p = doc.add_paragraph('Az alábbi dokumentum tartalmazza az NBA csapatok logóit és városaikat.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Gyűjtsük ki a csapatinformációkat a fájlnevekből és az NBA_TEAMS adatokból
    teams_info = {}
    
    # Először az NBA_TEAMS adatokat használjuk alapként
    for team in NBA_TEAMS:
        team_id = team["id"]
        teams_info[team_id] = {
            'name': team["name"],
            'abbreviation': team["abbreviation"],
            'city': team["city"],
            'logo_path': None  # Először nincs logó, majd a fájlnevek alapján beállítjuk
        }
    
    # Rendeljük hozzá a logó fájlokat a megfelelő csapatokhoz
    for logo_path in logo_files:
        file_name = os.path.basename(logo_path)
        # Várható formátum: ID_RÖVIDÍTÉS_NÉV.png vagy ID_RÖVIDÍTÉS_NÉV.svg
        for team_id, team_info in teams_info.items():
            abbr = team_info['abbreviation']
            if f"{team_id}_{abbr}_" in file_name:
                teams_info[team_id]['logo_path'] = logo_path
                break
    
    # Ha még mindig vannak csapatok logó nélkül, készítsünk helyettesítő képeket
    for team_id, team_info in teams_info.items():
        if not team_info['logo_path']:
            team_name = team_info['name']
            team_abbr = team_info['abbreviation']
            team_city = team_info['city']
            
            print(f"Hiányzó logó pótlása: {team_city} {team_name}")
            placeholder_png = os.path.join(logos_dir, f"{team_id}_{team_abbr}_{team_name}.png")
            create_text_image(f"{team_city}\n{team_name}", placeholder_png)
            teams_info[team_id]['logo_path'] = placeholder_png
    
    # Rendezzük a csapatokat konferenciánként (Keleti és Nyugati)
    eastern_teams = []
    western_teams = []
    
    # Keleti konferencia csapatai
    eastern_abbrs = ["BOS", "BKN", "NYK", "PHI", "TOR", "CHI", "CLE", "DET", "IND", "MIL", 
                      "ATL", "CHA", "MIA", "ORL", "WAS"]
    
    # Csapatok rendezése konferenciák szerint
    for team_id, team_info in teams_info.items():
        if team_info['abbreviation'] in eastern_abbrs:
            eastern_teams.append((team_id, team_info))
        else:
            western_teams.append((team_id, team_info))
    
    # Ellenőrizzük, hogy minden csapat szerepel-e
    all_abbrs = set([team_info['abbreviation'] for team_id, team_info in eastern_teams + western_teams])
    expected_abbrs = set([team["abbreviation"] for team in NBA_TEAMS])
    if all_abbrs != expected_abbrs:
        missing = expected_abbrs - all_abbrs
        if missing:
            print(f"FIGYELEM: Hiányzó csapatok: {missing}")
    
    # Dokumentum szélességének beállítása (hogy elférjen 4 oszlop)
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4 szélesség
    section.page_height = Cm(29.7)  # A4 magasság
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    
    # Minden csapat egy oldalon, 4 oszlopos elrendezésben
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    
    # Állítsuk be a táblázat oszlopainak szélességét
    for cell in table.columns[0].cells:
        cell.width = Cm(4.5)
    for cell in table.columns[1].cells:
        cell.width = Cm(4.5)
    for cell in table.columns[2].cells:
        cell.width = Cm(4.5)
    for cell in table.columns[3].cells:
        cell.width = Cm(4.5)
    
    # Minden csapat egy listába
    all_teams = []
    
    # Az Eastern Conference csapatai jönnek előre
    eastern_teams.sort(key=lambda x: x[1]['city'])  # Rendezzük városnév szerint
    for team in eastern_teams:
        all_teams.append(team)
    
    # Utána a Western Conference csapatai
    western_teams.sort(key=lambda x: x[1]['city'])  # Rendezzük városnév szerint
    for team in western_teams:
        all_teams.append(team)
    
    # Feltöltjük a táblázatot minden csapattal
    current_row = None
    
    for i, (team_id, team_info) in enumerate(all_teams):
        if i % 4 == 0:
            current_row = table.add_row()
            current_row.height = Cm(2.5)  # Kisebb sormagasság a 8 sor érdekében
        
        cell = current_row.cells[i % 4]
        team_name = team_info.get('name', '')
        team_city = team_info.get('city', '')
        team_abbr = team_info.get('abbreviation', '')
        logo_path = team_info.get('logo_path')
        
        # Külön debug üzenet az utolsó sor csapataihoz
        if team_abbr in last_row_abbrs:
            print(f"Utolsó sor csapata: {team_city} {team_name}, logó: {os.path.basename(logo_path) if logo_path else 'Nincs'}")
        
        if logo_path and os.path.exists(logo_path):
            # Hozzáadjuk a képet és a nevet
            p = cell.paragraphs[0]
            run = p.add_run()
            
            try:
                # Beillesztjük a képet, még kisebb méretben a 8 sor érdekében
                run.add_picture(logo_path, width=Cm(2.0))
            except Exception as e:
                print(f"Hiba a kép beillesztésekor ({team_name}): {str(e)}")
                run.add_text(f"[{team_city} {team_name}]")
            
            # Adjuk hozzá a csapat nevét és városát
            name_text = f"{team_city} {team_name}"
            
            # Név hozzáadása a cellához
            name_paragraph = cell.add_paragraph(name_text)
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Beállítjuk a szöveg betűméretét 9-re
            for run in name_paragraph.runs:
                run.font.size = Pt(9)
            
            # Csökkentsük a bekezdések közötti térközt
            name_paragraph.paragraph_format.space_after = Pt(2)
            
        else:
            # Ha nincs kép, csak a nevet és várost adjuk hozzá
            p = cell.paragraphs[0]
            p.text = f"{team_city} {team_name}"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Beállítjuk a szöveg betűméretét 9-re
            for run in p.runs:
                run.font.size = Pt(9)
    
    # Mentés
    try:
        doc.save(output_file)
        print(f"A Word dokumentum sikeresen létrehozva: {output_file}")
        return True
    except Exception as e:
        print(f"Hiba a Word dokumentum mentésekor: {str(e)}")
        traceback.print_exc()
        return False

def main():
    # Paraméterek ellenőrzése
    if len(sys.argv) > 1 and sys.argv[1] == '--download-only':
        # Csak logók letöltése
        output_dir = sys.argv[2] if len(sys.argv) > 2 else "nba_logos"
        download_nba_logos(output_dir)
        return
    
    # Alapértelmezett mappa és fájlnév
    logos_dir = "nba_logos"
    output_file = "nba_teams_logos.docx"
    
    # Ellenőrizzük, hogy létezik-e a logók mappa, ha nem, akkor létrehozzuk és letöltjük a logókat
    if not os.path.exists(logos_dir) or len(os.listdir(logos_dir)) == 0:
        print(f"A {logos_dir} mappa nem létezik vagy üres. Letöltjük a logókat...")
        download_nba_logos(logos_dir)
    
    # Word dokumentum létrehozása
    result = create_nba_teams_word_document(logos_dir, output_file)
    
    if result:
        print(f"A dokumentum elkészült: {output_file}")
    else:
        print("A dokumentum létrehozása sikertelen.")

if __name__ == "__main__":
    main() 