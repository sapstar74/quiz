import os
import json
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import traceback

# Próbáljuk betölteni az NHL csapatokat tartalmazó adatokat
try:
    from nhl_logos_update import NHL_TEAMS
    NHL_TEAMS_LOADED = True
except ImportError:
    NHL_TEAMS_LOADED = False
    print("Figyelmeztetés: Az NHL_TEAMS adatstruktúra nem érhető el. Csak a letöltött képek alapján dolgozunk.")

# Próbáljuk importálni a CairoSVG-t az SVG konvertáláshoz
try:
    import cairosvg
    CAIROSVG_AVAILABLE = True
    print("CairoSVG sikeresen betöltve. Az SVG-k konvertálása lehetséges.")
except ImportError:
    CAIROSVG_AVAILABLE = False
    print("CairoSVG nem elérhető. Az SVG képek nem lesznek konvertálva PNG-re.")

# A városok és csapatnevek definíciója
NHL_CITIES = {
    "NJD": "New Jersey",
    "NYI": "New York",
    "NYR": "New York",
    "PHI": "Philadelphia",
    "PIT": "Pittsburgh",
    "BOS": "Boston",
    "BUF": "Buffalo",
    "MTL": "Montréal",
    "OTT": "Ottawa",
    "TOR": "Toronto",
    "CAR": "Carolina", 
    "FLA": "Florida",
    "TBL": "Tampa Bay",
    "WSH": "Washington",
    "CHI": "Chicago",
    "DET": "Detroit",
    "NSH": "Nashville",
    "STL": "St. Louis",
    "CGY": "Calgary",
    "COL": "Colorado",
    "EDM": "Edmonton",
    "VAN": "Vancouver",
    "ANA": "Anaheim",
    "DAL": "Dallas",
    "LAK": "Los Angeles",
    "SJS": "San Jose",
    "CBJ": "Columbus",
    "MIN": "Minnesota",
    "WPG": "Winnipeg",
    "ARI": "Arizona",
    "VGK": "Vegas",
    "SEA": "Seattle"
}

def convert_svg_to_png(svg_path, png_path):
    """
    SVG fájl konvertálása PNG-re (ha lehetséges)
    """
    try:
        if CAIROSVG_AVAILABLE:
            # CairoSVG használata
            cairosvg.svg2png(url=svg_path, write_to=png_path)
            return True
        else:
            # Inkscape parancssoros hívása alternatívaként
            inkscape_cmd = f"inkscape --export-filename={png_path} {svg_path}"
            if os.system(inkscape_cmd) == 0 and os.path.exists(png_path):
                return True
    except Exception as e:
        print(f"Hiba az SVG->PNG konvertálás során: {str(e)}")
    
    return False

def create_text_image(text, output_path, width=200, height=200):
    """
    Egyszerű szöveges kép létrehozása
    """
    try:
        # Létrehozunk egy üres képet
        img = PILImage.new('RGB', (width, height), color=(240, 240, 240))
        
        # Szöveg kirajzolásához betöltjük az ImageDraw modult
        from PIL import ImageDraw, ImageFont
        draw = ImageDraw.Draw(img)
        
        # Alapértelmezett betűtípus
        font = ImageFont.load_default()
        
        # Szöveg elhelyezése középre
        lines = text.split("\n")
        y_pos = (height - len(lines) * 20) // 2
        
        for line in lines:
            # Szöveghossz becslése (pontos mérés nem mindig elérhető)
            textwidth = font.getlength(line) if hasattr(font, 'getlength') else len(line) * 6
            x_pos = (width - textwidth) // 2
            
            # Szöveg kirajzolása
            draw.text((x_pos, y_pos), line, fill=(0, 0, 0))
            y_pos += 20
        
        # Kép mentése
        img.save(output_path)
        return True
    except Exception as e:
        print(f"Hiba a szöveges kép létrehozása közben: {str(e)}")
        return False

def get_city_name(team_abbr):
    """
    Városi név lekérése a csapat rövidítése alapján
    """
    if team_abbr in NHL_CITIES:
        return NHL_CITIES[team_abbr]
    return None

def create_nhl_teams_word_document(logos_dir="nhl_logos_new", output_file="nhl_teams_logos.docx"):
    """
    Word dokumentum létrehozása NHL csapatok logóival
    """
    print(f"Word dokumentum létrehozása a következő mappából: {logos_dir}")
    
    # Ellenőrizzük, hogy létezik-e a mappa
    if not os.path.exists(logos_dir):
        print(f"A megadott mappa nem létezik: {logos_dir}")
        return False
    
    # Nézzük meg, van-e letöltött információnk a logókról
    results_file = os.path.join(logos_dir, "download_results.json")
    results_data = {}
    
    if os.path.exists(results_file):
        try:
            with open(results_file, 'r', encoding='utf-8') as f:
                results_data = json.load(f)
            print(f"Sikeresen betöltöttük a letöltési eredményeket: {len(results_data.get('successful_downloads', {}))} csapat")
        except Exception as e:
            print(f"Hiba a letöltési eredmények betöltésekor: {str(e)}")
    
    # SVG képek konvertálása PNG-re (Word nem kezeli az SVG-ket)
    svg_files = []
    for root, _, files in os.walk(logos_dir):
        for file in files:
            if file.lower().endswith('.svg'):
                svg_path = os.path.join(root, file)
                png_path = svg_path.replace('.svg', '.png')
                
                if not os.path.exists(png_path):
                    print(f"SVG konvertálása PNG-re: {file}")
                    if not convert_svg_to_png(svg_path, png_path):
                        # Ha a konvertálás sikertelen, készítsünk egy szöveges képet helyette
                        team_name = " ".join(file.split('_')[1:]).replace('.svg', '')
                        print(f"  Helyettesítő kép készítése: {team_name}")
                        create_text_image(team_name, png_path)
    
    # Gyűjtsük össze az elérhető PNG/JPG képeket
    logo_files = []
    for root, _, files in os.walk(logos_dir):
        for file in files:
            if file.lower().endswith(('.png', '.jpg', '.jpeg')):
                logo_files.append(os.path.join(root, file))
    
    if not logo_files:
        print(f"Nem találtunk képfájlokat a következő mappában: {logos_dir}")
        return False
    
    print(f"Összesen {len(logo_files)} képet találtunk")
    
    # Hozzuk létre az új Word dokumentumot
    doc = Document()
    
    # Állítsuk be a dokumentum tulajdonságait
    doc.core_properties.title = "NHL Csapatok Logói"
    doc.core_properties.author = "NHL Logo Downloader"
    
    # Címlap
    title = doc.add_heading('NHL Csapatok', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Rövid magyarázat
    p = doc.add_paragraph('Az alábbi dokumentum tartalmazza az NHL csapatok logóit és városaikat.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Gyűjtsük ki a csapatinformációkat a fájlnevekből
    teams_info = {}
    for logo_path in logo_files:
        file_name = os.path.basename(logo_path)
        # Várható formátum: ID_RÖVIDÍTÉS_NÉV.png
        parts = file_name.split('_', 2)  # Maximum 2 felosztás, a többit hagyjuk a névnek
        
        if len(parts) >= 2:
            try:
                team_id = int(parts[0])
                team_abbr = parts[1] if len(parts) > 1 else ""
                
                if len(parts) > 2:
                    team_name = parts[2].rsplit('.', 1)[0].replace('_', ' ')
                else:
                    team_name = parts[1].rsplit('.', 1)[0].replace('_', ' ')
                
                # Próbáljuk lekérni a város nevét
                city_name = get_city_name(team_abbr)
                
                teams_info[team_id] = {
                    'name': team_name,
                    'abbreviation': team_abbr,
                    'city': city_name,
                    'logo_path': logo_path
                }
            except (ValueError, IndexError):
                # Ha nem sikerül az ID kinyerése, akkor fájlnév alapján listázzuk
                pass
    
    # Ha rendelkezésre áll az NHL_TEAMS és a letöltési információ, használjuk azt is
    if NHL_TEAMS_LOADED:
        # Kiegészítjük a teams_info-t az NHL_TEAMS adataival
        for team in NHL_TEAMS:
            team_id = team.get('id')
            team_abbr = team.get('abbreviation', '')
            
            if team_id not in teams_info:
                # Városi név lekérése
                city_name = get_city_name(team_abbr)
                
                teams_info[team_id] = {
                    'name': team.get('name', ''),
                    'abbreviation': team_abbr,
                    'city': city_name,
                    'logo_path': None  # Nincs logó
                }
            else:
                # Frissítjük a meglévő adatokat
                teams_info[team_id]['abbreviation'] = team_abbr
                
                # Hozzáadjuk a város nevét, ha nem létezik még
                if 'city' not in teams_info[team_id] or not teams_info[team_id]['city']:
                    teams_info[team_id]['city'] = get_city_name(team_abbr)
                
                if 'name' not in teams_info[team_id] or not teams_info[team_id]['name']:
                    teams_info[team_id]['name'] = team.get('name', '')
    
    # Egyszerű listázási mód, 4 oszlopos táblázatban
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    
    # Feltöltjük a táblázatot
    current_row = None
    
    # Rendezzük a csapatokat ID szerint
    sorted_teams = sorted(teams_info.items())
    
    for i, (team_id, team_info) in enumerate(sorted_teams):
        if i % 4 == 0:
            current_row = table.add_row()
        
        cell = current_row.cells[i % 4]
        team_name = team_info.get('name', f'Csapat {team_id}')
        team_city = team_info.get('city', '')
        logo_path = team_info.get('logo_path')
        
        if logo_path and os.path.exists(logo_path):
            # Hozzáadjuk a képet és a nevet
            p = cell.paragraphs[0]
            run = p.add_run()
            
            try:
                # Beillesztjük a képet
                run.add_picture(logo_path, width=Cm(3))
            except Exception as e:
                print(f"Hiba a kép beillesztésekor ({team_name}): {str(e)}")
                run.add_text("[Kép Hiányzik]")
            
            # Adjuk hozzá a csapat nevét és városát
            if team_city:
                name_text = f"{team_city} {team_name}"
            else:
                name_text = team_name
            
            cell.add_paragraph(name_text).alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Ha nincs kép, csak a nevet és várost adjuk hozzá
            if team_city:
                cell.text = f"{team_city} {team_name}"
            else:
                cell.text = team_name
    
    # Mentés
    try:
        doc.save(output_file)
        print(f"A Word dokumentum sikeresen létrehozva: {output_file}")
        return True
    except Exception as e:
        print(f"Hiba a Word dokumentum mentésekor: {str(e)}")
        traceback.print_exc()
        return False

def main():
    # Paraméterek ellenőrzése
    if len(sys.argv) > 1:
        logos_dir = sys.argv[1]
    else:
        logos_dir = "nhl_logos_new"
    
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        output_file = "nhl_teams_logos.docx"
    
    # Word dokumentum létrehozása
    result = create_nhl_teams_word_document(logos_dir, output_file)
    
    if result:
        print(f"A dokumentum elkészült: {output_file}")
    else:
        print("A dokumentum létrehozása sikertelen.")

if __name__ == "__main__":
    main() 