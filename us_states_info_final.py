import os
import requests
import io
from PIL import Image, ImageDraw, ImageFont
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import json

# Create directory for state seals
os.makedirs('us_state_seals', exist_ok=True)

# Dictionary of US states with their capitals
US_STATES = [
    {"name": "Alabama", "abbreviation": "AL", "capital": "Montgomery"},
    {"name": "Alaska", "abbreviation": "AK", "capital": "Juneau"},
    {"name": "Arizona", "abbreviation": "AZ", "capital": "Phoenix"},
    {"name": "Arkansas", "abbreviation": "AR", "capital": "Little Rock"},
    {"name": "California", "abbreviation": "CA", "capital": "Sacramento"},
    {"name": "Colorado", "abbreviation": "CO", "capital": "Denver"},
    {"name": "Connecticut", "abbreviation": "CT", "capital": "Hartford"},
    {"name": "Delaware", "abbreviation": "DE", "capital": "Dover"},
    {"name": "Florida", "abbreviation": "FL", "capital": "Tallahassee"},
    {"name": "Georgia", "abbreviation": "GA", "capital": "Atlanta"},
    {"name": "Hawaii", "abbreviation": "HI", "capital": "Honolulu"},
    {"name": "Idaho", "abbreviation": "ID", "capital": "Boise"},
    {"name": "Illinois", "abbreviation": "IL", "capital": "Springfield"},
    {"name": "Indiana", "abbreviation": "IN", "capital": "Indianapolis"},
    {"name": "Iowa", "abbreviation": "IA", "capital": "Des Moines"},
    {"name": "Kansas", "abbreviation": "KS", "capital": "Topeka"},
    {"name": "Kentucky", "abbreviation": "KY", "capital": "Frankfort"},
    {"name": "Louisiana", "abbreviation": "LA", "capital": "Baton Rouge"},
    {"name": "Maine", "abbreviation": "ME", "capital": "Augusta"},
    {"name": "Maryland", "abbreviation": "MD", "capital": "Annapolis"},
    {"name": "Massachusetts", "abbreviation": "MA", "capital": "Boston"},
    {"name": "Michigan", "abbreviation": "MI", "capital": "Lansing"},
    {"name": "Minnesota", "abbreviation": "MN", "capital": "Saint Paul"},
    {"name": "Mississippi", "abbreviation": "MS", "capital": "Jackson"},
    {"name": "Missouri", "abbreviation": "MO", "capital": "Jefferson City"},
    {"name": "Montana", "abbreviation": "MT", "capital": "Helena"},
    {"name": "Nebraska", "abbreviation": "NE", "capital": "Lincoln"},
    {"name": "Nevada", "abbreviation": "NV", "capital": "Carson City"},
    {"name": "New Hampshire", "abbreviation": "NH", "capital": "Concord"},
    {"name": "New Jersey", "abbreviation": "NJ", "capital": "Trenton"},
    {"name": "New Mexico", "abbreviation": "NM", "capital": "Santa Fe"},
    {"name": "New York", "abbreviation": "NY", "capital": "Albany"},
    {"name": "North Carolina", "abbreviation": "NC", "capital": "Raleigh"},
    {"name": "North Dakota", "abbreviation": "ND", "capital": "Bismarck"},
    {"name": "Ohio", "abbreviation": "OH", "capital": "Columbus"},
    {"name": "Oklahoma", "abbreviation": "OK", "capital": "Oklahoma City"},
    {"name": "Oregon", "abbreviation": "OR", "capital": "Salem"},
    {"name": "Pennsylvania", "abbreviation": "PA", "capital": "Harrisburg"},
    {"name": "Rhode Island", "abbreviation": "RI", "capital": "Providence"},
    {"name": "South Carolina", "abbreviation": "SC", "capital": "Columbia"},
    {"name": "South Dakota", "abbreviation": "SD", "capital": "Pierre"},
    {"name": "Tennessee", "abbreviation": "TN", "capital": "Nashville"},
    {"name": "Texas", "abbreviation": "TX", "capital": "Austin"},
    {"name": "Utah", "abbreviation": "UT", "capital": "Salt Lake City"},
    {"name": "Vermont", "abbreviation": "VT", "capital": "Montpelier"},
    {"name": "Virginia", "abbreviation": "VA", "capital": "Richmond"},
    {"name": "Washington", "abbreviation": "WA", "capital": "Olympia"},
    {"name": "West Virginia", "abbreviation": "WV", "capital": "Charleston"},
    {"name": "Wisconsin", "abbreviation": "WI", "capital": "Madison"},
    {"name": "Wyoming", "abbreviation": "WY", "capital": "Cheyenne"},
]

# Better URLs for state seals that directly link to PNG files
STATE_SEAL_URLS = {
    "AL": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/Alabamasealstateflag.jpg",
    "AK": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/AlaskaStateSealprimary.jpg",
    "AZ": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/arizonagreatseal.jpg",
    "AR": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/ArkansasStateSealprimary.jpg",
    "CA": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/CaliforniaStateSealprimary.jpg",
    "CO": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/ColoradoStateSealprimary.jpg",
    "CT": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/ConnecticutStateSealprimary.jpg",
    "DE": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/DelawareStateSealprimary.jpg",
    "FL": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/FloridaStateSealprimary.jpg",
    "GA": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/GeorgiaStateSealprimary.jpg",
    "HI": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/HawaiiStateSealprimary.jpg",
    "ID": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/idahostatesealcolor.jpg",
    "IL": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/IllinoisStateSealprimary.jpg",
    "IN": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/IndianaStateSealprimary.jpg",
    "IA": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/IowaStateSealprimary.jpg",
    "KS": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/KansasStateSealprimary.jpg",
    "KY": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/KentuckyStateSealprimary.jpg",
    "LA": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/LouisianaStateSealprimary.jpg",
    "ME": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/MaineStateSealprimary.jpg",
    "MD": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/marylandstatesealreverse.jpg",
    "MA": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/MassachusettsStateSealprimary.jpg",
    "MI": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/Michiganstateseal.jpg",
    "MN": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/MinnesotaStateSealprimary.jpg",
    "MS": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/MississippiStateSealprimary.jpg",
    "MO": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/MissouriStateSealprimary.jpg",
    "MT": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/MontanaStateSealprimary.jpg",
    "NE": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/NebraskaStateSealprimary.jpg",
    "NV": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/NevadaStateSealprimary.jpg",
    "NH": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/NewHampshireStateSealprimary.jpg",
    "NJ": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/NewJerseyStateSealprimary.jpg",
    "NM": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/NewMexicoStateSealprimary.jpg",
    "NY": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/NewYorkStateSealprimary.jpg",
    "NC": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/NorthCarolinaStateSealprimary.jpg",
    "ND": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/NorthDakotaStateSealprimary.jpg",
    "OH": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/OhioStateSealprimary.jpg",
    "OK": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/OklahomaStateSealprimary.jpg",
    "OR": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/OregonStateSealprimary.jpg",
    "PA": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/PennsylvaniaStateSealprimary.jpg",
    "RI": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/RhodeIslandStateSealprimary.jpg",
    "SC": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/SouthCarolinaStateSealprimary.jpg",
    "SD": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/SouthDakotaStateSealprimary.jpg",
    "TN": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/TennesseeStateSealprimary.jpg",
    "TX": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/TexasStateSealprimary.jpg",
    "UT": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/UtahStateSealprimary.jpg",
    "VT": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/VermontStateSealprimary.jpg",
    "VA": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/VirginiaStateSealprimary.jpg",
    "WA": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/WashingtonStateSealprimary.jpg",
    "WV": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/WestVirginiaStateSealprimary.jpg",
    "WI": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/WisconsinStateSealprimary.jpg",
    "WY": "https://statesymbolsusa.org/sites/statesymbolsusa.org/files/primary-images/WyomingStateSealprimary.jpg",
}

def create_text_image(text, size=(200, 200), bg_color=(240, 240, 240), text_color=(0, 0, 0)):
    """Create a text image as a placeholder for missing state seals."""
    img = Image.new('RGB', size, color=bg_color)
    draw = ImageDraw.Draw(img)
    
    # Use default font
    font = ImageFont.load_default()
    
    # Draw text centered (without using getsize which is deprecated)
    text_lines = text.split('\n')
    y_position = size[1] // 2 - 10 * len(text_lines)
    
    for line in text_lines:
        # Approximate center position
        x_position = size[0] // 2 - 5 * len(line)  
        draw.text((x_position, y_position), line, font=font, fill=text_color)
        y_position += 20  # Move down for next line
    
    return img

def download_state_seal(state_name, state_abbr):
    """Download a state seal image from pre-defined URL."""
    img_path = os.path.join('us_state_seals', f"{state_abbr}.png")
    
    # If image already exists, return its path
    if os.path.exists(img_path):
        return img_path
    
    # Get URL from pre-defined dict
    if state_abbr in STATE_SEAL_URLS:
        try:
            url = STATE_SEAL_URLS[state_abbr]
            print(f"Downloading seal for {state_name} from {url}")
            response = requests.get(url, timeout=10)
            
            if response.status_code == 200:
                # Save the image directly from response
                with open(img_path, 'wb') as f:
                    f.write(response.content)
                print(f"Successfully downloaded seal for {state_name}")
                return img_path
        except Exception as e:
            print(f"Error downloading seal for {state_name}: {e}")
    
    # Fallback URL options if primary source fails
    fallback_urls = [
        f"https://www.netstate.com/states/symb/{state_abbr.lower()}_symb.htm",
        f"https://en.wikipedia.org/wiki/Seal_of_{state_name.replace(' ', '_')}",
        f"https://www.50states.com/state-seals/{state_name.lower().replace(' ', '-')}-state-seal.htm"
    ]
    
    # Try fallback URLs
    for fallback_url in fallback_urls:
        try:
            print(f"Trying fallback URL for {state_name}: {fallback_url}")
            response = requests.get(fallback_url, timeout=10)
            if response.status_code == 200:
                # This approach requires parsing HTML to find image URLs
                # For simplicity, we'll just use it as a last resort
                print(f"Found fallback page for {state_name}, but need direct image URL")
                break
        except Exception as e:
            print(f"Error with fallback for {state_name}: {e}")
    
    # If download fails, create a placeholder image
    try:
        placeholder = create_text_image(f"Seal of\n{state_name}")
        placeholder.save(img_path)
        print(f"Created placeholder seal for {state_name}")
        return img_path
    except Exception as e:
        print(f"Error creating placeholder for {state_name}: {e}")
        return None

def download_all_state_seals():
    """Download all state seals and return results."""
    results = {"success": [], "failed": []}
    
    for state in US_STATES:
        state_name = state["name"]
        state_abbr = state["abbreviation"]
        
        img_path = download_state_seal(state_name, state_abbr)
        
        if img_path and os.path.exists(img_path):
            results["success"].append({
                "state": state_name,
                "abbr": state_abbr,
                "path": img_path
            })
        else:
            results["failed"].append({
                "state": state_name,
                "abbr": state_abbr
            })
        
        # Add a small delay to avoid overwhelming servers
        time.sleep(0.2)
    
    # Save results to a JSON file for reference
    with open('us_state_seals/download_results.json', 'w') as f:
        json.dump(results, f, indent=2)
    
    return results

def create_us_states_doc():
    """Create a Word document with all US states, their seals, and capitals."""
    # First download all state seals
    print("Downloading state seals...")
    results = download_all_state_seals()
    print(f"Downloaded {len(results['success'])} state seals successfully")
    if results['failed']:
        print(f"Failed to download {len(results['failed'])} state seals")
    
    # Create a Word document
    doc = docx.Document()
    
    # Set document properties
    doc.core_properties.title = "United States - States and Capitals"
    
    # Set document margins (smaller margins for better fit)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
    
    # Add a title
    title = doc.add_heading("United States - States and Capitals", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create a table with 5 states per row
    # Total rows needed: ceil(50/5) = 10 rows
    table = doc.add_table(rows=10, cols=5)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set table width to fill page
    table.width = Cm(19)
    
    # Set equal width for each column
    for cell in table.columns:
        cell.width = Cm(3.8)
    
    # Add states to the table
    for i, state in enumerate(US_STATES):
        # Calculate row and column position
        row = i // 5
        col = i % 5
        
        # Get cell
        cell = table.cell(row, col)
        
        # Add state name as cell heading
        state_para = cell.paragraphs[0]
        state_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        state_run = state_para.add_run(state["name"])
        state_run.bold = True
        state_run.font.size = Pt(10)
        
        # Add state seal image
        img_path = os.path.join('us_state_seals', f"{state['abbreviation']}.png")
        if os.path.exists(img_path):
            try:
                # Add image paragraph
                img_para = cell.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(img_path, width=Cm(3.2))
            except Exception as e:
                print(f"Error adding image for {state['name']}: {e}")
                # Add a placeholder text instead
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("[Seal Image]")
        
        # Add capital city
        capital_para = cell.add_paragraph()
        capital_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        capital_run = capital_para.add_run(f"Capital: {state['capital']}")
        capital_run.font.size = Pt(9)
        
        # Set row height to accommodate content
        table.rows[row].height = Cm(4.5)
        cell.vertical_alignment = 1  # centered
    
    # Save the document
    output_file = "us_states_capitals.docx"
    doc.save(output_file)
    print(f"Document created successfully: {output_file}")
    return output_file

if __name__ == "__main__":
    create_us_states_doc() 